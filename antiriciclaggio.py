import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext, filedialog
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import Dict, List, Optional
import json
import os
import sys

# =====================================================================
# GESTIONE CONFIGURAZIONI JSON
# =====================================================================

class ConfigLoader:
    """Gestisce il caricamento delle configurazioni JSON esterne."""
    
    def __init__(self, config_dir: str = "config"):
        self.config_dir = config_dir
        self.configs = {}
        self._load_all_configs()
    
    def _get_config_path(self, filename: str) -> str:
        """Restituisce il path completo del file di configurazione."""
        if hasattr(sys, '_MEIPASS'):
            base_path = sys._MEIPASS
        else:
            base_path = os.path.dirname(os.path.abspath(__file__))
        
        return os.path.join(base_path, self.config_dir, filename)
    
    def _load_json(self, filename: str) -> dict:
        """Carica un file JSON con gestione errori."""
        filepath = self._get_config_path(filename)
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
                print(f"✓ Caricato: {filename}")
                return data
        except FileNotFoundError:
            print(f"✗ ERRORE: File {filename} non trovato in {filepath}")
            messagebox.showerror(
                "Errore Configurazione",
                f"File di configurazione mancante: {filename}\n\n"
                f"Posizione attesa: {filepath}\n\n"
                f"Assicurati che la cartella 'config' contenga tutti i file JSON necessari."
            )
            sys.exit(1)
        except json.JSONDecodeError as e:
            print(f"✗ ERRORE: File {filename} non è un JSON valido: {e}")
            messagebox.showerror(
                "Errore Configurazione",
                f"Errore nel file {filename}:\n{str(e)}\n\n"
                f"Il file JSON contiene errori di sintassi."
            )
            sys.exit(1)
        except Exception as e:
            print(f"✗ ERRORE: Impossibile caricare {filename}: {e}")
            messagebox.showerror(
                "Errore Configurazione",
                f"Errore imprevisto caricando {filename}:\n{str(e)}"
            )
            sys.exit(1)
    
    def _load_all_configs(self):
        """Carica tutti i file di configurazione necessari."""
        print("\n" + "="*60)
        print("CARICAMENTO CONFIGURAZIONI JSON")
        print("="*60)
        
        self.configs['luoghi'] = self._load_json('luoghi_rischio.json')
        self.configs['natura'] = self._load_json('natura_giuridica.json')
        self.configs['prestazioni'] = self._load_json('prestazioni_veda.json')
        self.configs['fattori'] = self._load_json('fattori_rischio.json')
        self.configs['generale'] = self._load_json('configurazione.json')
        self.configs['clienti'] = self._load_json('clienti_studio.json')
        self.configs['avvocati'] = self._load_json('avvocati.json')
        
        print("="*60)
        print("✓ TUTTI I FILE CARICATI CORRETTAMENTE")
        print("="*60 + "\n")
    
    def get(self, config_name: str, *keys):
        """Recupera un valore dalla configurazione usando chiavi annidate."""
        try:
            value = self.configs[config_name]
            for key in keys:
                value = value[key]
            return value
        except (KeyError, TypeError):
            print(f"⚠ Chiave non trovata: {config_name} -> {' -> '.join(keys)}")
            return None

_config_loader = None

def get_config() -> ConfigLoader:
    """Restituisce l'istanza globale del ConfigLoader."""
    global _config_loader
    if _config_loader is None:
        _config_loader = ConfigLoader()
    return _config_loader

# =====================================================================
# CONFIGURAZIONE
# =====================================================================

class Config:
    """Contiene tutte le configurazioni dell'applicazione caricate da JSON."""
    
    @staticmethod
    def get_sections_A_manual() -> List[Dict]:
        loader = get_config()
        tabella_a = loader.get('fattori', 'tabella_a_aspetti_cliente')
        sections = []
        for nome, dati in tabella_a.items():
            if not nome.startswith('_'):
                sections.append({"name": nome, "factors": dati["fattori"]})
        return sections
    
    @staticmethod
    def get_sections_B_manual() -> List[Dict]:
        loader = get_config()
        tabella_b = loader.get('fattori', 'tabella_b_aspetti_operazione')
        sections = []
        for nome, dati in tabella_b.items():
            if not nome.startswith('_'):
                sections.append({"name": nome, "factors": dati["fattori"]})
        return sections
    
    @staticmethod
    def get_natura_giuridica_database() -> Dict:
        loader = get_config()
        db_esatto = loader.get('natura', 'database_esatto')
        result = {}
        for categoria, items in db_esatto.items():
            if not categoria.startswith('_'):
                if isinstance(items, dict):
                    livello = items.get('_livello', 3)
                    for nome, descrizione in items.items():
                        if not nome.startswith('_'):
                            result[nome] = (livello, descrizione)
        return result
    
    @staticmethod
    def get_natura_giuridica_keywords() -> Dict:
        loader = get_config()
        keywords_data = loader.get('natura', 'keywords_ricerca')
        return {
            1: keywords_data.get('livello_1', []),
            2: keywords_data.get('livello_2', []),
            3: keywords_data.get('livello_3', []),
            4: keywords_data.get('livello_4', [])
        }
    
    @staticmethod
    def get_natura_giuridica_categorie_manuali() -> Dict:
        loader = get_config()
        return loader.get('natura', 'categorie_manuali') or {}
    
    @staticmethod
    def get_clienti_studio() -> List[str]:
        """Restituisce la lista semplice dei clienti studio (senza livello)."""
        loader = get_config()
        return loader.get('clienti', 'clienti') or []
    
    @staticmethod
    def get_avvocati_studio() -> List[str]:
        """Restituisce la lista degli avvocati dello studio."""
        loader = get_config()
        avvocati_dict = loader.get('avvocati', 'avvocati') or {}
        return list(avvocati_dict.keys())
    
    @staticmethod
    def get_province_italiane_rischio() -> Dict:
        loader = get_config()
        province = loader.get('luoghi', 'province_italiane')
        return {k: v for k, v in province.items() if not k.startswith('_')}
    
    @staticmethod
    def get_paesi_rischio() -> Dict:
        loader = get_config()
        paesi = loader.get('luoghi', 'paesi_internazionali')
        return {k: v for k, v in paesi.items() if not k.startswith('_')}
    
    LUOGHI_RISCHIO_COMPLETO = {}
    
    @classmethod
    def inizializza_luoghi(cls):
        """Crea il dizionario unificato di tutti i luoghi."""
        cls.LUOGHI_RISCHIO_COMPLETO = {}
        cls.LUOGHI_RISCHIO_COMPLETO.update(cls.get_province_italiane_rischio())
        cls.LUOGHI_RISCHIO_COMPLETO.update(cls.get_paesi_rischio())
        print(f"✓ Inizializzati {len(cls.LUOGHI_RISCHIO_COMPLETO)} luoghi")
    
    @staticmethod
    def get_prestazioni_veda() -> Dict[str, int]:
        loader = get_config()
        prestazioni_data = loader.get('prestazioni', 'prestazioni')
        result = {}
        for nome, dati in prestazioni_data.items():
            result[nome] = dati.get('livello', 3)
        return result
    
    @staticmethod
    def get_prestazioni_solo_tabella_a() -> List[str]:
        loader = get_config()
        return loader.get('prestazioni', 'prestazioni_solo_tabella_a', 'lista') or []
    
    @staticmethod
    def get_soglia_importo() -> float:
        loader = get_config()
        return loader.get('generale', 'normativa', 'soglia_adeguata_verifica', 'valore') or 15000.0
    
    @staticmethod
    def get_color_scheme() -> Dict:
        loader = get_config()
        colori = loader.get('generale', 'interfaccia_grafica', 'colori')
        return {
            "bg": colori.get('background', '#000000'),
            "frame_bg": colori.get('frame_background', '#ECF0F1'),
            "input_bg": colori.get('input_background', '#000000'),
            "text_fg": colori.get('text_foreground', '#2C3E50'),
            "button_bg": colori.get('button_background', '#2C3E50'),
            "button_fg": colori.get('button_foreground', '#000000'),
            "button_hover": colori.get('button_hover', '#34495E'),
            "button_active": colori.get('button_active', '#1A252F'),
            "select_color": colori.get('selection_color', '#D3E0EA'),
            "tooltip_bg": colori.get('tooltip_background', '#FFFEF0'),
            "tooltip_fg": colori.get('tooltip_foreground', '#000000'),
            "auto_bg": colori.get('auto_section_background', '#E8F8F5'),
            "auto_label": colori.get('auto_section_label', '#27AE60')
        }
    
    RISCHIO_INERENTE_OPTIONS = [
        "1 - Non significativo",
        "2 - Poco significativo",
        "3 - Abbastanza significativo",
        "4 - Molto significativo"
    ]

# =====================================================================
# TOOLTIP
# =====================================================================

class ToolTip:
    """Gestisce un singolo tooltip per un widget specifico."""
    
    def __init__(self, widget: tk.Widget, text: str, delay: int = 500):
        self.widget = widget
        self.text = text
        self.delay = delay
        self.tooltip_window = None
        self.schedule_id = None
        self.widget.bind('<Enter>', self._on_enter)
        self.widget.bind('<Leave>', self._on_leave)
        self.widget.bind('<Button>', self._on_leave)
    
    def _on_enter(self, event=None):
        self._cancel_scheduled()
        self.schedule_id = self.widget.after(self.delay, self._show_tooltip)
    
    def _on_leave(self, event=None):
        self._cancel_scheduled()
        self._hide_tooltip()
    
    def _cancel_scheduled(self):
        if self.schedule_id:
            self.widget.after_cancel(self.schedule_id)
            self.schedule_id = None
    
    def _show_tooltip(self):
        if self.tooltip_window:
            return
        color_scheme = Config.get_color_scheme()
        x = self.widget.winfo_rootx() + 10
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 5
        self.tooltip_window = tk.Toplevel(self.widget)
        self.tooltip_window.wm_overrideredirect(True)
        self.tooltip_window.wm_geometry(f"+{x}+{y}")
        label = tk.Label(self.tooltip_window, text=self.text, 
                        background=color_scheme["tooltip_bg"], 
                        foreground=color_scheme["tooltip_fg"], 
                        relief="solid", borderwidth=2, padx=12, pady=8, 
                        wraplength=500, justify="left", font=("Helvetica", 11, "bold"))
        label.pack()
    
    def _hide_tooltip(self):
        if self.tooltip_window:
            self.tooltip_window.destroy()
            self.tooltip_window = None

# =====================================================================
# CALCOLATORE RISCHIO
# =====================================================================

class RiskCalculator:
    """Gestisce i calcoli del rischio secondo metodologia VEDA."""
    
    @staticmethod
    def calcola_media_sezione(section_data: Dict) -> float:
        selected_levels = [section_data["level_vars"][i].get() 
                          for i, var in enumerate(section_data["vars"]) if var.get()]
        return round(sum(selected_levels) / len(selected_levels), 1) if selected_levels else 0.0
    
    @staticmethod
    def calcola_totale_sezioni(sections: List[Dict], auto_sections: List[Dict] = None) -> tuple:
        totale = 0.0
        sub_medias = []
        if auto_sections:
            for sec in auto_sections:
                media = sec["level"]
                sub_medias.append(media)
                totale += media
        for sec in sections:
            media = RiskCalculator.calcola_media_sezione(sec)
            sub_medias.append(media)
            totale += media
        return totale, sub_medias
    
    @staticmethod
    def calcola_livello_rischio(somma: float) -> str:
        if somma <= 2.5:
            return "BASSO/POCO SIGNIFICATIVO - Adeguata Verifica Semplificata"
        elif somma <= 3.5:
            return "ABBASTANZA SIGNIFICATIVO - Adeguata Verifica Ordinaria"
        else:
            return "MOLTO SIGNIFICATIVO - Adeguata Verifica Rafforzata"
    
    @staticmethod
    def valida_anomalia_importo(importo: Optional[float] = None) -> str:
        soglia = Config.get_soglia_importo()
        if importo is None:
            return f"Importo non specificato - Nessuna anomalia per soglia €{soglia:,.2f}"
     #   elif importo >= soglia:
     #       return f"ATTENZIONE: Importo ≥ €{soglia:,.2f} - ADEGUATA VERIFICA OBBLIGATORIA"
     #   else:
     #       return f"Importo < €{soglia:,.2f} - Nessun obbligo di adeguata verifica per importo"
    
    @staticmethod
    def calcola_livello_da_importo(importo: float) -> int:
        if importo < 50000:
            return 1
        elif importo < 250000:
            return 2
        elif importo < 1000000:
            return 3
        else:
            return 4
    
    @staticmethod
    def get_livello_paese(paese: str) -> int:
        paese_norm = paese.strip().title()
        return Config.LUOGHI_RISCHIO_COMPLETO.get(paese_norm, 0)
    
    @staticmethod
    def cerca_natura_giuridica(testo: str) -> list:
        testo_lower = testo.lower().strip()
        if len(testo_lower) < 2:
            return []
        
        risultati = []
        database = Config.get_natura_giuridica_database()
        keywords = Config.get_natura_giuridica_keywords()
        clienti_studio = Config.get_clienti_studio()
        
        # 1. Database standard (con livello)
        for nome, (livello, descrizione) in database.items():
            if testo_lower in nome.lower():
                risultati.append((nome, livello, descrizione, False))
        
        # 2. Clienti studio (senza livello)
        for cliente in clienti_studio:
            if testo_lower in cliente.lower():
                risultati.append((cliente, None, "Cliente studio", True))
        
        # 3. Keywords fallback
        if not risultati:
            for livello, keyword_list in keywords.items():
                for keyword in keyword_list:
                    if keyword in testo_lower:
                        tipo_generico = {
                            1: "Soggetto vigilato / Grande azienda",
                            2: "Media impresa / Studio professionale",
                            3: "PMI / Piccola società",
                            4: "Alto rischio / Trust / Offshore"
                        }
                        risultati.append((f"Categoria: {tipo_generico[livello]}", 
                                        livello, tipo_generico[livello], False))
                        break
                if risultati:
                    break
        
        risultati.sort(key=lambda x: (x[3], x[1] if x[1] else 999, x[0]))
        return risultati[:10]
    
    @staticmethod
    def cerca_luogo(testo: str) -> list:
        testo_lower = testo.lower().strip()
        if len(testo_lower) < 2:
            return []
        risultati = []
        for luogo, livello in Config.LUOGHI_RISCHIO_COMPLETO.items():
            if testo_lower in luogo.lower():
                risultati.append((luogo, livello))
        risultati.sort(key=lambda x: (not x[0].lower().startswith(testo_lower), x[1], x[0]))
        return risultati[:10]

# =====================================================================
# WORD EXPORTER
# =====================================================================

class WordExporter:
    """Gestisce l'export dei dati in formato Word."""
    
    @staticmethod
    def esporta(dati: Dict, sections_def_A: List[Dict], sections_def_B: List[Dict]) -> None:
        doc = Document()
        WordExporter._configura_stili(doc)
        title = doc.add_paragraph()
        title_run = title.add_run(f"VALUTAZIONE DEL RISCHIO AML - {dati['cliente']}")
        title_run.bold = True
        title_run.font.size = Pt(18)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Data: {dati['data_valutazione']}")
        date_run.font.size = Pt(10)
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.add_paragraph()
        
        if dati.get("usa_solo_tabella_a", False):
            nota_veda = doc.add_paragraph()
            nota_veda.add_run("NOTA METODOLOGICA: ").bold = True
            nota_veda.add_run(f"La prestazione '{dati['prestazione_veda']}' richiede valutazione solo Tabella A")
            nota_veda.runs[0].font.color.rgb = RGBColor(255, 140, 0)
            doc.add_paragraph()
        
        WordExporter._aggiungi_titolo(doc, "I. Misurazione del rischio specifico")
        WordExporter._aggiungi_sezione_mista(doc, "A. Aspetti connessi al cliente", 
                                             dati["sections_A_manual"], dati["sections_A_auto"],
                                             sections_def_A, dati["total_A"])
        doc.add_paragraph()
        
        if not dati.get("usa_solo_tabella_a", False):
            WordExporter._aggiungi_sezione_mista(doc, "B. Aspetti connessi all'operazione e/o prestazione professionale", 
                                                 dati["sections_B_manual"], dati["sections_B_auto"],
                                                 sections_def_B, dati["total_B"])
        else:
            nota_b = doc.add_paragraph()
            nota_b.add_run("B. Aspetti connessi all'operazione e/o prestazione professionale").bold = True
            doc.add_paragraph("TABELLA B NON COMPILATA - Prestazione continuativa (Linee Guida VEDA pag. 97)")

        # Tutto inline, nessun page break
        doc.add_paragraph()
        WordExporter._aggiungi_legenda_punteggi(doc)
        doc.add_paragraph()
        WordExporter._aggiungi_intestazione(doc, dati)
        WordExporter._aggiungi_sintesi_risultati(doc, dati)
        
        # Tronca nome cliente se troppo lungo per evitare problemi filesystem
        cliente_short = dati['cliente'][:30].replace(" ", "_")
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialdir="~",
            initialfile=f"{cliente_short}_Valutazione Rischio {datetime.now().strftime('%Y_%m_%d')}.docx",
            filetypes=[("Word files", "*.docx")]
        )
        if file_path:
            doc.save(file_path)
            messagebox.showinfo("Successo", f"File Word creato: {file_path}")
    
    @staticmethod
    def _configura_stili(doc: Document) -> None:
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Inches(0.1)
        style.paragraph_format.line_spacing = 1.15
    
    @staticmethod
    def _aggiungi_intestazione(doc: Document, dati: Dict) -> None:
        title = doc.add_paragraph()
        title_run = title.add_run("PROFILATURA DEL RISCHIO AML")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        info_table = doc.add_table(rows=8, cols=2)
        info_table.style = 'Light Grid Accent 1'

        # Imposta larghezza colonne
        info_table.columns[0].width = Cm(6)   # Etichetta
        info_table.columns[1].width = Cm(11)  # Valore

        info_data = [
            ("Cliente:", dati["cliente"]), 
            ("Titolare Effettivo:", dati["titolare_effettivo"]), 
            ("Scopo Operazione:", dati["scopo_operazione"]), 
            ("Descrizione Attività:", dati["descrizione_attivita"]), 
            ("Importo Operazione:", f"€ {dati['importo']:,.2f}" if dati['importo'] else "Non specificato"), 
            ("Prestazione Professionale:", dati.get("prestazione_veda", "Non specificata")),
            ("Avvocato:", dati["avvocato"]), 
            ("Rischio Inerente:", f"{dati['rischio_inerente']} - {Config.RISCHIO_INERENTE_OPTIONS[dati['rischio_inerente']-1].split(' - ')[1]}")
        ]
        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
            for cell in row.cells:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(11)
                    cell.paragraphs[0].runs[0].font.name = 'Arial'
            if row.cells[0].paragraphs[0].runs:
                row.cells[0].paragraphs[0].runs[0].bold = True
        doc.add_paragraph()
    
    @staticmethod
    def _aggiungi_sintesi_risultati(doc: Document, dati: Dict) -> None:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        doc.add_page_break()
        sintesi_title = doc.add_paragraph()
        sintesi_title_run = sintesi_title.add_run("SINTESI DELLA VALUTAZIONE E CALCOLI")
        sintesi_title_run.bold = True
        sintesi_title_run.font.size = Pt(16)
        sintesi_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        calc_title = doc.add_paragraph()
        calc_title_run = calc_title.add_run("DETTAGLIO DEI CALCOLI")
        calc_title_run.bold = True
        calc_title_run.font.size = Pt(13)
        calc_title_run.underline = True
        
        num_rows = 7 if dati.get("usa_solo_tabella_a", False) else 8
        calc_table = doc.add_table(rows=num_rows, cols=2)
        calc_table.style = 'Light Grid Accent 1'

        # Imposta larghezza colonne
        calc_table.columns[0].width = Cm(11)  # Descrizione
        calc_table.columns[1].width = Cm(6)   # Valore

        if dati.get("usa_solo_tabella_a", False):
            calc_data = [
                ("1. Totale A (Aspetti Cliente)", f"{dati['total_A']:.2f}"),
                ("", ""),
                ("2. Rischio Specifico = A / numero fattori", f"{dati['total_A']:.2f} / {dati['num_fattori_a']} = {dati['rischio_specifico']:.2f}"),
                ("", ""),
                ("3. Rischio Inerente (VEDA)", f"{dati['rischio_inerente']}"),
                ("", ""),
                ("4. RISULTATO FINALE", f"Somma Ponderata = {dati['somma']:.2f} - {dati['livello']}")
            ]
        else:
            calc_data = [
                ("1. Totale A (Aspetti Cliente)", f"{dati['total_A']:.2f}"),
                ("2. Totale B (Aspetti Operazione)", f"{dati['total_B']:.2f}"),
                ("", ""),
                ("3. Rischio Specifico = (A + B) / 10", f"({dati['total_A']:.2f} + {dati['total_B']:.2f}) / 10 = {dati['rischio_specifico']:.2f}"),
                ("", ""),
                ("4. Rischio Inerente (VEDA)", f"{dati['rischio_inerente']}"),
                ("", ""),
                ("5. RISULTATO FINALE", f"Somma Ponderata = {dati['somma']:.2f} - {dati['livello']}")
            ]
        
        for i, (label, value) in enumerate(calc_data):
            row = calc_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
            for cell in row.cells:
                if cell.text and cell.paragraphs[0].runs:
                    para = cell.paragraphs[0]
                    para.runs[0].font.size = Pt(11)
                    para.runs[0].font.name = 'Arial'
                    if "RISULTATO FINALE" in label or "Rischio Specifico" in label:
                        para.runs[0].bold = True
                        para.runs[0].font.size = Pt(12)
        doc.add_paragraph()
        firma_para = doc.add_paragraph()
        firma_para.add_run(f"Data: {dati['data_valutazione']}\n\n")
        firma_para.add_run(f"Avvocato: {dati['avvocato']}\n\n")
        firma_para.add_run("Firma: _______________________________")
        for run in firma_para.runs:
            run.font.size = Pt(11)
        # formula_para = doc.add_paragraph()
        # formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # formula_text = f"Somma Ponderata = (Inerente × 0.3) + (Specifico × 0.7)\nSomma Ponderata = ({dati['rischio_inerente']} × 0.3) + ({dati['rischio_specifico']:.2f} × 0.7)\nSomma Ponderata = {dati['inerente_ponderato']:.2f} + {dati['specifico_ponderato']:.2f} = {dati['somma']:.2f}"
        # formula_run = formula_para.add_run(formula_text)
        # formula_run.font.size = Pt(12)
        # formula_run.bold = True
        # formula_run.font.name = 'Courier New'
        # pPr = formula_para._element.get_or_add_pPr()
        # shading_elm = OxmlElement('w:shd')
        # shading_elm.set(qn('w:fill'), 'E8F4F8')
        # pPr.append(shading_elm)
        # doc.add_paragraph()
        # doc.add_paragraph()
        # result_title = doc.add_paragraph("RISULTATO FINALE")
        # result_title_run = result_title.runs[0]
        # result_title_run.bold = True
        # result_title_run.font.size = Pt(14)
        # result_title_run.underline = True
        # result_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # result_table = doc.add_table(rows=2, cols=1)
        # result_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # row0 = result_table.rows[0]
        # row0.cells[0].text = f"SOMMA PONDERATA: {dati['somma']:.2f}"
        # para0 = row0.cells[0].paragraphs[0]
        # para0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # para0.runs[0].font.size = Pt(14)
        # para0.runs[0].bold = True
        # row1 = result_table.rows[1]
        # row1.cells[0].text = f"LIVELLO DI RISCHIO: {dati['livello']}"
        # para1 = row1.cells[0].paragraphs[0]
        # para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # para1.runs[0].font.size = Pt(16)
        # para1.runs[0].bold = True
        # livello = dati['livello']
        # if "NON SIGNIFICATIVO" in livello:
        #     color_hex = '90EE90'
        #     para1.runs[0].font.color.rgb = RGBColor(0, 100, 0)
        # elif "POCO SIGNIFICATIVO" in livello:
        #     color_hex = 'FFFF00'
        #     para1.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        # elif "ABBASTANZA SIGNIFICATIVO" in livello:
        #     color_hex = 'FFA500'
        #     para1.runs[0].font.color.rgb = RGBColor(139, 0, 0)
        # else:
        #     color_hex = 'FF6B6B'
        #     para1.runs[0].font.color.rgb = RGBColor(139, 0, 0)
        # shading = OxmlElement('w:shd')
        # shading.set(qn('w:fill'), color_hex)
        # row1.cells[0]._element.get_or_add_tcPr().append(shading)
        # doc.add_paragraph()

        # Sezione Rischio Rafforzato (se RE > 3.5)
        if dati.get("dichiarazione_rafforzata"):
            WordExporter._aggiungi_sezione_rischio_rafforzato(doc, dati)

        anomalie_para = doc.add_paragraph()
        run_label = anomalie_para.add_run("")
        run_label.bold = True
        run_label.font.size = Pt(12)
        run_text = anomalie_para.add_run(dati['anomalie'])
        run_text.font.size = Pt(11)
        # Niente colori - testo nero su bianco
        doc.add_paragraph()

    @staticmethod
    def _aggiungi_legenda_punteggi(doc: Document) -> None:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        legenda_title = doc.add_paragraph()
        legenda_title_run = legenda_title.add_run("LEGENDA DEI LIVELLI DI RISCHIO")
        legenda_title_run.bold = True
        legenda_title_run.font.size = Pt(12)

        # Tabella compatta 4 righe x 2 colonne
        legenda_table = doc.add_table(rows=4, cols=2)
        legenda_table.style = 'Table Grid'

        # Imposta larghezza colonne - compatta
        legenda_table.columns[0].width = Cm(4)   # Range
        legenda_table.columns[1].width = Cm(8)   # Descrizione

        # Dati legenda con colori
        legenda_data = [
            ("≤ 2.5", "NON SIGNIFICATIVO", "90EE90", RGBColor(0, 100, 0)),
            ("2.5 < x ≤ 3.0", "POCO SIGNIFICATIVO", "FFFF00", RGBColor(0, 0, 0)),
            ("3.0 < x ≤ 3.5", "ABBASTANZA SIGNIFICATIVO", "FFA500", RGBColor(139, 0, 0)),
            ("> 3.5", "MOLTO SIGNIFICATIVO", "FF6B6B", RGBColor(139, 0, 0))
        ]

        for i, (range_val, descrizione, bg_color, text_color) in enumerate(legenda_data):
            row = legenda_table.rows[i]

            # Cella range - usa add_run invece di text
            cell0 = row.cells[0]
            para0 = cell0.paragraphs[0]
            run0 = para0.add_run(range_val)
            run0.font.size = Pt(10)
            run0.bold = True
            para0.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Cella descrizione - usa add_run invece di text
            cell1 = row.cells[1]
            para1 = cell1.paragraphs[0]
            run1 = para1.add_run(descrizione)
            run1.font.size = Pt(10)
            run1.bold = True
            run1.font.color.rgb = text_color

            # Sfondo colorato per entrambe le celle
            for cell in row.cells:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), bg_color)
                cell._element.get_or_add_tcPr().append(shading)
    
    @staticmethod
    def _aggiungi_sezione_rischio_rafforzato(doc: Document, dati: Dict) -> None:
        """Aggiunge la sezione di dichiarazione per Rischio Rafforzato (RE > 3.5)"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        dichiarazione = dati["dichiarazione_rafforzata"]

        doc.add_paragraph()

        # DICHIARAZIONE DEL PROFESSIONISTA
        dich_title = doc.add_paragraph()
        run_title = dich_title.add_run("DICHIARAZIONE DEL PROFESSIONISTA")
        run_title.bold = True
        run_title.font.size = Pt(12)

        # Testo dichiarazione
        dich_text = f"""Avvocato {dati['avvocato']} DICHIARA di aver preso visione degli obblighi di Adeguata Verifica Rafforzata (D.Lgs. 231/2007) e si impegna ad adempierli separatamente, conservando la relativa documentazione nel fascicolo del cliente {dati['cliente']}."""

        dich_para = doc.add_paragraph()
        run_dich = dich_para.add_run(dich_text)
        run_dich.font.size = Pt(11)

        doc.add_paragraph()

        # Data, Nome e Firma
        firma_para = doc.add_paragraph()
        firma_para.add_run(f"Data: {dati['data_valutazione']}\n\n")
        firma_para.add_run(f"Avvocato: {dati['avvocato']}\n\n")
        firma_para.add_run("Firma: _______________________________")
        for run in firma_para.runs:
            run.font.size = Pt(11)

    @staticmethod
    def _aggiungi_titolo(doc: Document, testo: str) -> None:
        title = doc.add_paragraph(testo)
        title_run = title.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(14)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    @staticmethod
    def _estrai_numero_fattore(nome: str) -> float:
        """Estrae il numero dal nome del fattore (es. 'A.1' -> 1.0, 'B.12' -> 12.0)"""
        import re
        match = re.search(r'[A-Z]\.(\d+(?:\.\d+)?)', nome)
        return float(match.group(1)) if match else 999.0

    @staticmethod
    def _aggiungi_sezione_mista(doc: Document, titolo: str, section_data_manual: List[Dict], 
                               section_data_auto: List[Dict], section_def: List[Dict], totale: float) -> None:
        doc.add_paragraph(titolo, style='Normal')
        doc.add_paragraph("Fattore di rischio riscontrato", style='Normal')
        doc.add_paragraph("Livello di rischio specifico (da 1 a 4)", style='Normal')
        table = doc.add_table(rows=1, cols=3)
        table.style = 'TableGrid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Imposta larghezza colonne
        COL_WIDTH_FATTORE = Cm(13.5)
        COL_WIDTH_X = Cm(0.8)
        COL_WIDTH_LIVELLO = Cm(1.5)

        table.columns[0].width = COL_WIDTH_FATTORE
        table.columns[1].width = COL_WIDTH_X
        table.columns[2].width = COL_WIDTH_LIVELLO

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Fattore"
        hdr_cells[1].text = ""
        hdr_cells[2].text = "Livello"
        hdr_cells[0].width = COL_WIDTH_FATTORE
        hdr_cells[1].width = COL_WIDTH_X
        hdr_cells[2].width = COL_WIDTH_LIVELLO
        WordExporter._formatta_cella_custom(hdr_cells[0], 14, 'left', bold=True, bg_color='E8E8E8')
        WordExporter._formatta_cella_custom(hdr_cells[1], 14, 'center', bold=True, bg_color='E8E8E8')
        WordExporter._formatta_cella_custom(hdr_cells[2], 14, 'center', bold=True, bg_color='E8E8E8')

        # Crea lista unificata con tipo di sezione
        all_sections = []
        for sec in section_data_auto:
            all_sections.append({'type': 'auto', 'data': sec})
        for sec in section_data_manual:
            all_sections.append({'type': 'manual', 'data': sec})

        # Ordina tutti i fattori per numero
        all_sections_sorted = sorted(all_sections, key=lambda x: WordExporter._estrai_numero_fattore(x['data']['name']))

        # Itera sulla lista ordinata con zebrato
        row_index = 1  # Inizia da 1 perché 0 è l'header
        for item in all_sections_sorted:
            if item['type'] == 'auto':
                sec = item['data']
                bg = 'F5F5F5' if row_index % 2 == 0 else None
                row = table.add_row().cells
                row[0].text = f"{sec['name']}"
                row[1].text = "x"
                row[2].text = str(sec["level"])
                row[0].width = COL_WIDTH_FATTORE
                row[1].width = COL_WIDTH_X
                row[2].width = COL_WIDTH_LIVELLO
                WordExporter._formatta_cella_custom(row[0], 11, 'left', bold=True, bg_color=bg)
                WordExporter._formatta_cella_custom(row[1], 11, 'center', bg_color=bg)
                WordExporter._formatta_cella_custom(row[2], 11, 'center', bg_color=bg)
                row_index += 1

                bg = 'F5F5F5' if row_index % 2 == 0 else None
                row = table.add_row().cells
                row[0].text = sec["value"]
                row[0].width = COL_WIDTH_FATTORE
                row[1].width = COL_WIDTH_X
                row[2].width = COL_WIDTH_LIVELLO
                WordExporter._formatta_cella_custom(row[0], 10, 'left', bg_color=bg)
                WordExporter._formatta_cella_custom(row[1], 10, 'left', bg_color=bg)
                WordExporter._formatta_cella_custom(row[2], 10, 'left', bg_color=bg)
                row_index += 1
            else:  # manual
                sec = item['data']
                bg = 'F5F5F5' if row_index % 2 == 0 else None
                row = table.add_row().cells
                row[0].text = sec["name"]
                row[0].width = COL_WIDTH_FATTORE
                row[1].width = COL_WIDTH_X
                row[2].width = COL_WIDTH_LIVELLO
                WordExporter._formatta_cella_custom(row[0], 11, 'left', bold=True, bg_color=bg)
                WordExporter._formatta_cella_custom(row[1], 11, 'center', bg_color=bg)
                WordExporter._formatta_cella_custom(row[2], 11, 'center', bg_color=bg)
                row_index += 1

                for i, factor in enumerate(sec["factors"]):
                    bg = 'F5F5F5' if row_index % 2 == 0 else None
                    row = table.add_row().cells
                    row[0].text = factor
                    row[0].width = COL_WIDTH_FATTORE
                    row[1].width = COL_WIDTH_X
                    row[2].width = COL_WIDTH_LIVELLO
                    if sec["vars"][i].get():
                        row[1].text = "x"
                        row[2].text = str(sec["level_vars"][i].get())
                    WordExporter._formatta_cella_custom(row[0], 11, 'left', bg_color=bg)
                    WordExporter._formatta_cella_custom(row[1], 11, 'center', bg_color=bg)
                    WordExporter._formatta_cella_custom(row[2], 11, 'center', bg_color=bg)
                    row_index += 1
        
        row = table.add_row().cells
        row[0].text = f"TOTALE {titolo[0]}"
        row[2].text = str(round(totale, 1))
        row[0].width = COL_WIDTH_FATTORE
        row[1].width = COL_WIDTH_X
        row[2].width = COL_WIDTH_LIVELLO
        WordExporter._formatta_cella_custom(row[0], 14, 'left', bold=True)
        WordExporter._formatta_cella_custom(row[2], 14, 'center', bold=True)
    
    @staticmethod
    def _formatta_cella_custom(cell, font_size: int, align: str, bold: bool = False, bg_color: str = None) -> None:
        # Imposta sfondo anche se la cella è vuota
        if bg_color:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), bg_color)
            cell._element.get_or_add_tcPr().append(shading_elm)

        if not cell.text:
            return

        paragraph = cell.paragraphs[0]
        if align == 'left':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if paragraph.runs:
            run = paragraph.runs[0]
            run.font.size = Pt(font_size)
            run.font.name = 'Arial'
            run.bold = bold
    
    @staticmethod
    def _aggiungi_note_pep(doc: Document) -> None:
        """Metodo placeholder per note PEP (opzionale)"""
        pass

# =====================================================================
# APPLICAZIONE PRINCIPALE
# =====================================================================

class AMLRiskApp:
    """Applicazione principale per la profilatura del rischio AML."""
    
    def __init__(self, root: tk.Tk):
        self.root = root
        self.root.title("Profilatura Rischio AML")
        Config.inizializza_luoghi()
        
        color_scheme = Config.get_color_scheme()
        MIN_WIDTH = 1200
        MIN_HEIGHT = 700
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        window_width = max(int(screen_width * 0.90), MIN_WIDTH)
        window_height = max(int(screen_height * 0.85), MIN_HEIGHT)
        x_position = (screen_width - window_width) // 2
        y_position = (screen_height - window_height) // 2
        self.root.geometry(f"{window_width}x{window_height}+{x_position}+{y_position}")
        self.root.minsize(MIN_WIDTH, MIN_HEIGHT)
        self.root.configure(bg=color_scheme["bg"])
        
        self.section_vars_A_manual = []
        self.section_vars_B_manual = []
        self.entry_natura_giuridica = None
        self.label_natura_livello = None
        self.combo_natura_fallback = None
        self.suggerimenti_natura = None
        self.a1_manual_factors = None  # MODIFICA 1 VEDA: Fattori soggettivi A.1
        self.entry_area_cliente = None
        self.label_livello_cliente = None
        self.entry_area_destinazione = None
        self.label_livello_destinazione = None
        self.label_importo_livello = None
        self.check_b3_incongruo = None
        self.check_b3_frazionamenti = None
        self.check_b3_altro = None
        self.label_b3_livello_finale = None
        self.label_b3_base = None
        self.suggerimenti_cliente = None
        self.suggerimenti_destinazione = None
        self.suggerimenti_cliente_nome = None
        self.suggerimenti_avvocato_nome = None
        self.entry_data = None
        self.dati_export = None
        self._setup_ui()
    
    def _setup_ui(self) -> None:
        canvas, scrollable_frame = self._create_scrollable_canvas()
        self._create_general_inputs(scrollable_frame)
        self._create_sections(scrollable_frame)
        self._create_action_buttons(scrollable_frame)
        color_scheme = Config.get_color_scheme()
        self.label_risultato = tk.Label(scrollable_frame, text="", justify="left", wraplength=700, 
                                       bg=color_scheme["bg"], fg=color_scheme["text_fg"], 
                                       font=("Helvetica", 10))
        self.label_risultato.pack(pady=10)
    
    def _create_scrollable_canvas(self) -> tuple:
        color_scheme = Config.get_color_scheme()
        container = tk.Frame(self.root, bg=color_scheme["bg"])
        container.pack(fill="both", expand=True)
        canvas = tk.Canvas(container, bg=color_scheme["bg"])
        v_scrollbar = tk.Scrollbar(container, orient="vertical", command=canvas.yview, 
                                   bg=color_scheme["button_bg"])
        h_scrollbar = tk.Scrollbar(container, orient="horizontal", command=canvas.xview, 
                                   bg=color_scheme["button_bg"])
        scrollable_frame = tk.Frame(canvas, bg=color_scheme["bg"])
        canvas.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)
        
        def scroll_canvas(event):
            direction = -1 if (event.num == 4 or event.delta > 0) else 1
            canvas.yview_scroll(direction, "units")
        
        for event in ("<MouseWheel>", "<Button-4>", "<Button-5>"):
            canvas.bind_all(event, scroll_canvas)
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        
        def on_frame_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
        
        scrollable_frame.bind("<Configure>", on_frame_configure)
        canvas.grid(row=0, column=0, sticky="nsew")
        v_scrollbar.grid(row=0, column=1, sticky="ns")
        h_scrollbar.grid(row=1, column=0, sticky="ew")
        container.grid_rowconfigure(0, weight=1)
        container.grid_columnconfigure(0, weight=1)
        return canvas, scrollable_frame
    
    def _create_general_inputs(self, parent: tk.Frame) -> None:
        color_scheme = Config.get_color_scheme()
        title = tk.Label(parent, text="PROFILATURA RISCHIO ANTIRICICLAGGIO (AML)", 
                        font=("Helvetica", 16, "bold"), bg=color_scheme["bg"], 
                        fg=color_scheme["text_fg"])
        title.pack(pady=10)
        
        # subtitle = tk.Label(parent, text="v3.2 - Conforme Linee Guida VEDA 02-2020 + D.Lgs. 231/2007 - FIXED", 
        #                    font=("Helvetica", 9, "italic"), bg=color_scheme["bg"], 
        #                    fg="#7F8C8D")
        # subtitle.pack()
        
        main_container = tk.Frame(parent, bg=color_scheme["bg"])
        main_container.pack(fill="x", padx=20, pady=10)
        
        left_column = tk.Frame(main_container, bg=color_scheme["bg"])
        left_column.pack(side="left", fill="both", expand=True, padx=(0, 10))
        
        # DATA VALUTAZIONE
        tk.Label(left_column, text="Data Valutazione (gg/mm/aaaa):", 
                bg=color_scheme["bg"], fg=color_scheme["text_fg"], font=("Helvetica", 10, "bold")
               ).pack(anchor="w")

        self.entry_data = tk.Entry(left_column, bg=color_scheme["input_bg"], fg="#000000", 
                           borderwidth=2, width=15, font=("Helvetica", 10))
        self.entry_data.pack(anchor="w", pady=(0, 10))

        # Imposta data corrente come default
        self.entry_data.insert(0, datetime.now().strftime('%d/%m/%Y'))

        
        # NOME CLIENTE CON AUTOCOMPLETE
        tk.Label(left_column, text="Nome Cliente:", bg=color_scheme["bg"], 
                fg=color_scheme["text_fg"], font=("Helvetica", 10, "bold")).pack(anchor="w")
        
        cliente_container = tk.Frame(left_column, bg=color_scheme["bg"])
        cliente_container.pack(fill="x", pady=(0, 10))
        
        self.entry_cliente = tk.Entry(cliente_container, bg=color_scheme["input_bg"],
                                     fg="#000000", borderwidth=2, width=40, font=("Helvetica", 10))
        self.entry_cliente.pack(fill="x")
        
        self.suggerimenti_cliente_nome = tk.Listbox(cliente_container, height=6, bg="#FFFFFF", 
                                                    fg="#000000", font=("Helvetica", 9), 
                                                    selectbackground=color_scheme["select_color"])
        
        self.entry_cliente.bind('<KeyRelease>', self._on_cliente_nome_keyrelease)
        self.suggerimenti_cliente_nome.bind('<<ListboxSelect>>', self._on_cliente_nome_select)
        # 🔧 FIX: Chiude menù quando si clicca fuori dal campo
        self.entry_cliente.bind('<FocusOut>', lambda e: self.suggerimenti_cliente_nome.pack_forget())
        
        tk.Label(left_column, text="Titolare Effettivo:", bg=color_scheme["bg"], 
                fg=color_scheme["text_fg"], font=("Helvetica", 10, "bold")).pack(anchor="w")
        self.entry_titolare = tk.Entry(left_column, bg=color_scheme["input_bg"], 
                                      fg="#000000", borderwidth=2, width=40, font=("Helvetica", 10))
        self.entry_titolare.pack(fill="x", pady=(0, 10))
        
        tk.Label(left_column, text="Scopo Operazione:", bg=color_scheme["bg"], 
                fg=color_scheme["text_fg"], font=("Helvetica", 10, "bold")).pack(anchor="w")
        self.text_scopo = scrolledtext.ScrolledText(left_column, bg=color_scheme["input_bg"], 
                                                    fg="#000000", borderwidth=2, height=5, 
                                                    width=40, font=("Helvetica", 10))
        self.text_scopo.pack(fill="both", expand=True)
        
        right_column = tk.Frame(main_container, bg=color_scheme["bg"])
        right_column.pack(side="right", fill="both", expand=True, padx=(10, 0))
        
        tk.Label(right_column, text="Descrizione Attività:", bg=color_scheme["bg"], 
                fg=color_scheme["text_fg"], font=("Helvetica", 10, "bold")).pack(anchor="w")
        self.text_attivita = scrolledtext.ScrolledText(right_column, bg=color_scheme["input_bg"], 
                                                       fg="#000000", borderwidth=2, height=5, 
                                                       width=40, font=("Helvetica", 10))
        self.text_attivita.pack(fill="both", expand=True, pady=(0, 10))
        
        tk.Label(right_column, text="Importo Operazione (€):", bg=color_scheme["bg"], 
                fg=color_scheme["text_fg"], font=("Helvetica", 10, "bold")).pack(anchor="w")
        self.entry_importo = tk.Entry(right_column, bg=color_scheme["input_bg"], 
                                     fg="#000000", borderwidth=2, width=40, font=("Helvetica", 10))
        self.entry_importo.pack(fill="x", pady=(0, 5))
        
        self.label_importo_livello = tk.Label(right_column, text="", 
                                             bg=color_scheme["auto_bg"], 
                                             fg=color_scheme["auto_label"], 
                                             font=("Helvetica", 9, "bold"), pady=3)
        
        self.entry_importo.bind('<KeyRelease>', self._aggiorna_livello_importo)
        
        # AVVOCATO CON AUTOCOMPLETE
        tk.Label(right_column, text="Avvocato:", bg=color_scheme["bg"], 
                fg=color_scheme["text_fg"], font=("Helvetica", 10, "bold")).pack(anchor="w")
        
        avvocato_container = tk.Frame(right_column, bg=color_scheme["bg"])
        avvocato_container.pack(fill="x", pady=(0, 10))
        
        self.entry_avvocato = tk.Entry(avvocato_container, bg=color_scheme["input_bg"],
                                      fg="#000000", borderwidth=2, width=40, font=("Helvetica", 10))
        self.entry_avvocato.pack(fill="x")
        
        self.suggerimenti_avvocato_nome = tk.Listbox(avvocato_container, height=6, bg="#FFFFFF", 
                                                     fg="#000000", font=("Helvetica", 9), 
                                                     selectbackground=color_scheme["select_color"])
        
        self.entry_avvocato.bind('<KeyRelease>', self._on_avvocato_nome_keyrelease)
        self.suggerimenti_avvocato_nome.bind('<<ListboxSelect>>', self._on_avvocato_nome_select)
        # 🔧 FIX: Chiude menù quando si clicca fuori dal campo
        self.entry_avvocato.bind('<FocusOut>', lambda e: self.suggerimenti_avvocato_nome.pack_forget())
        
        rischio_container = tk.Frame(parent, bg=color_scheme["auto_bg"], 
                                     borderwidth=2, relief="solid", padx=20, pady=15)
        rischio_container.pack(pady=15, padx=20, fill="x")
        
        title_frame = tk.Frame(rischio_container, bg=color_scheme["auto_bg"])
        title_frame.pack()
        
        tk.Label(title_frame, text="RISCHIO INERENTE", 
                bg=color_scheme["auto_bg"], fg=color_scheme["auto_label"], 
                font=("Helvetica", 11, "bold")).pack(side="left")
        
        info_label = tk.Label(title_frame, text=" ℹ️", bg=color_scheme["auto_bg"], 
                             font=("Helvetica", 12), cursor="question_arrow")
        info_label.pack(side="left", padx=5)
        
        tooltip_inerente = """RISCHIO INERENTE - TABELLA VEDA UFFICIALE:

Il livello di rischio viene assegnato AUTOMATICAMENTE
in base alla prestazione professionale svolta.

FONTE: Linee Guida VEDA Edizione 02-2020, pag. 99
Tabella vincolante CNDCEC dal 01/01/2020

LIVELLI:
- Livello 2 (POCO): Consulenza contrattuale, tributaria, 
  amministrazione liquidazione, custodia beni, valutazioni
  
- Livello 3 (ABBASTANZA): Consulenza societaria, aziendale,
  costituzione società, trust, tenuta contabilità, revisione
  
- Livello 4 (MOLTO): Operazioni di finanza straordinaria

NOTA IMPORTANTE:
Per "Tenuta contabilità" e "Revisione legale conti" 
verrà compilata SOLO la Tabella A (senza Tabella B)
come previsto da VEDA pag. 97."""
        
        ToolTip(info_label, tooltip_inerente, delay=300)
        
        tk.Label(rischio_container, text="Seleziona Prestazione Professionale:", 
                bg=color_scheme["auto_bg"], fg=color_scheme["text_fg"], 
                font=("Helvetica", 10, "bold")).pack(pady=(10, 5))
        
        prestazioni_veda = Config.get_prestazioni_veda()
        prestazioni_list = list(prestazioni_veda.keys())
        self.combo_prestazione = tk.StringVar(value=prestazioni_list[0])
        
        # 🔧 FIX DEFINITIVO: CUSTOM COMBOBOX con estetica professionale
        combo_container = tk.Frame(rischio_container, bg=color_scheme["auto_bg"])
        combo_container.pack(padx=10, pady=5, fill="x")
        
        # Frame per Entry + Freccia (simulazione combobox)
        combo_frame = tk.Frame(combo_container, bg="#FFFFFF", relief="solid", borderwidth=2)
        combo_frame.pack(fill="x")
        
        # Entry readonly che mostra la selezione
        self.combo_entry = tk.Entry(combo_frame, textvariable=self.combo_prestazione,
                                    bg="#FFFFFF", fg="#000000",
                                    font=("Helvetica", 10), 
                                    readonlybackground="#FFFFFF",
                                    state="readonly",
                                    cursor="hand2", relief="flat", borderwidth=0,
                                    disabledbackground="#FFFFFF",
                                    disabledforeground="#000000")
        self.combo_entry.pack(side="left", fill="both", expand=True, padx=8, pady=8)
        
        # Frame per la freccia con bordo sinistro
        arrow_frame = tk.Frame(combo_frame, bg="#E0E0E0", width=25, relief="flat")
        arrow_frame.pack(side="right", fill="y")
        arrow_frame.pack_propagate(False)
        
        # Bottone freccia per aprire/chiudere
        btn_dropdown = tk.Label(arrow_frame, text="▼", bg="#E0E0E0",
                               fg="#333333", font=("Helvetica", 10, "bold"), 
                               cursor="hand2")
        btn_dropdown.place(relx=0.5, rely=0.5, anchor="center")
        
        # Listbox nascosta con le opzioni (con scrollbar)
        listbox_frame = tk.Frame(rischio_container, bg="#FFFFFF", relief="solid", borderwidth=2)
        
        scrollbar = tk.Scrollbar(listbox_frame, orient="vertical")
        
        self.combo_listbox = tk.Listbox(listbox_frame, height=8,
                                        bg="#FFFFFF", fg="#000000",
                                        font=("Helvetica", 10),
                                        selectbackground="#3498DB",
                                        selectforeground="#FFFFFF",
                                        activestyle="none",
                                        relief="flat",
                                        borderwidth=0,
                                        highlightthickness=0,
                                        yscrollcommand=scrollbar.set)
        
        scrollbar.config(command=self.combo_listbox.yview)
        scrollbar.pack(side="right", fill="y")
        self.combo_listbox.pack(side="left", fill="both", expand=True)
        
        for prestazione in prestazioni_list:
            self.combo_listbox.insert(tk.END, prestazione)
        
        # 🎯 FUNZIONI PER GESTIRE APERTURA/CHIUSURA
        def toggle_dropdown(event=None):
            """Apre/chiude il dropdown"""
            if listbox_frame.winfo_ismapped():
                close_dropdown()
            else:
                open_dropdown()
        
        def open_dropdown():
            """Apre il dropdown e cambia freccia"""
            listbox_frame.pack(in_=combo_container, fill="x", pady=(2, 0))
            btn_dropdown.config(text="▲")  # Freccia su
            combo_frame.config(relief="solid", borderwidth=2, bg="#3498DB")
            
            # Seleziona item corrente
            current = self.combo_prestazione.get()
            try:
                idx = prestazioni_list.index(current)
                self.combo_listbox.selection_clear(0, tk.END)
                self.combo_listbox.selection_set(idx)
                self.combo_listbox.see(idx)
            except ValueError:
                pass
            self.combo_listbox.focus_set()
        
        def close_dropdown(event=None):
            """Chiude il dropdown e ripristina freccia"""
            listbox_frame.pack_forget()
            btn_dropdown.config(text="▼")  # Freccia giù
            combo_frame.config(relief="solid", borderwidth=2, bg="#FFFFFF")
        
        def select_item(event=None):
            """Seleziona un item e chiude il dropdown"""
            selection = self.combo_listbox.curselection()
            if selection:
                selected = self.combo_listbox.get(selection[0])
                self.combo_prestazione.set(selected)
                close_dropdown()
                self._aggiorna_rischio_inerente()
        
        def stop_scroll(event):
            """Blocca lo scroll della pagina quando cursore sul dropdown"""
            return "break"
        
        def on_hover_enter(event):
            """Evidenzia quando mouse passa sopra"""
            if not listbox_frame.winfo_ismapped():
                combo_frame.config(bg="#E8F4F8")
        
        def on_hover_leave(event):
            """Rimuove evidenziazione quando mouse esce"""
            if not listbox_frame.winfo_ismapped():
                combo_frame.config(bg="#FFFFFF")
        
        # 🔧 BIND per gestire interazioni
        self.combo_entry.bind('<Button-1>', toggle_dropdown)
        combo_frame.bind('<Button-1>', toggle_dropdown)
        btn_dropdown.bind('<Button-1>', toggle_dropdown)
        arrow_frame.bind('<Button-1>', toggle_dropdown)
        
        self.combo_listbox.bind('<<ListboxSelect>>', select_item)
        self.combo_listbox.bind('<Return>', select_item)
        self.combo_listbox.bind('<Double-Button-1>', select_item)
        self.combo_listbox.bind('<Escape>', close_dropdown)
        self.combo_listbox.bind('<FocusOut>', close_dropdown)
        
        # Hover effects
        combo_frame.bind('<Enter>', on_hover_enter)
        combo_frame.bind('<Leave>', on_hover_leave)
        
        # 🎯 QUESTO È IL FIX: Blocca scroll sul listbox!
        self.combo_listbox.bind("<MouseWheel>", stop_scroll)
        self.combo_listbox.bind("<Button-4>", stop_scroll)
        self.combo_listbox.bind("<Button-5>", stop_scroll)
        
        # Bind globale per chiudere con click fuori
        def check_click_outside(event):
            widget = event.widget
            # Se click fuori dal combo, chiudi
            if widget not in (self.combo_entry, btn_dropdown, self.combo_listbox, combo_frame, arrow_frame):
                if listbox_frame.winfo_ismapped():
                    close_dropdown()
        
        self.root.bind('<Button-1>', check_click_outside, add='+')
        
        self.label_rischio_inerente = tk.Label(rischio_container, text="", 
                                               bg=color_scheme["auto_bg"], 
                                               fg=color_scheme["auto_label"], 
                                               font=("Helvetica", 11, "bold"), pady=8)
        self.label_rischio_inerente.pack()
        
        self.label_solo_tabella_a = tk.Label(rischio_container, text="", 
                                             bg=color_scheme["auto_bg"], 
                                             fg="#E67E22", 
                                             font=("Helvetica", 9, "bold"), pady=5)
        self.label_solo_tabella_a.pack()
        
        # Aggiorna rischio inerente iniziale
        self._aggiorna_rischio_inerente()
        
        help_label = tk.Label(rischio_container, 
                             text="Il livello di rischio inerente è calcolato automaticamente", 
                             bg=color_scheme["auto_bg"], fg="#7F8C8D",
                             font=("Helvetica", 8, "italic"))
        help_label.pack(pady=(5, 0))
    
    def _on_cliente_nome_keyrelease(self, event=None) -> None:
        """Gestisce l'autocompletamento per il nome cliente."""
        # 🔧 FIX: ESC chiude il menù
        if event and event.keysym == 'Escape':
            self.suggerimenti_cliente_nome.pack_forget()
            return
        
        # 🔧 FIX: INVIO chiude il menù e accetta il testo
        if event and event.keysym == 'Return':
            self.suggerimenti_cliente_nome.pack_forget()
            return
        
        testo = self.entry_cliente.get()
        if len(testo) < 2:
            self.suggerimenti_cliente_nome.pack_forget()
            return
        
        clienti_studio = Config.get_clienti_studio()
        suggerimenti = []
        testo_lower = testo.lower()
        
        for cliente in clienti_studio:
            if testo_lower in cliente.lower():
                suggerimenti.append(cliente)
        
        if suggerimenti:
            self.suggerimenti_cliente_nome.delete(0, tk.END)
            for cliente in sorted(suggerimenti)[:10]:
                self.suggerimenti_cliente_nome.insert(tk.END, cliente)
            self.suggerimenti_cliente_nome.pack(fill="x", pady=(2, 0))
        else:
            self.suggerimenti_cliente_nome.pack_forget()
    
    def _on_cliente_nome_select(self, event=None) -> None:
        """Gestisce la selezione di un cliente dai suggerimenti."""
        try:
            selection = self.suggerimenti_cliente_nome.curselection()
            if selection:
                cliente_selezionato = self.suggerimenti_cliente_nome.get(selection[0])
                self.entry_cliente.delete(0, tk.END)
                self.entry_cliente.insert(0, cliente_selezionato)
                self.suggerimenti_cliente_nome.pack_forget()
                
                # BONUS: Auto-popola anche A.1 Natura Giuridica
                self.entry_natura_giuridica.delete(0, tk.END)
                self.entry_natura_giuridica.insert(0, cliente_selezionato)

                # Auto-rileva il livello di natura giuridica
                self._auto_detect_natura_giuridica(cliente_selezionato)
        except Exception as e:
            print(f"Errore selezione cliente: {e}")
    
    def _on_avvocato_nome_keyrelease(self, event=None) -> None:
        """Gestisce l'autocompletamento per il nome avvocato."""
        # 🔧 FIX: ESC chiude il menù
        if event and event.keysym == 'Escape':
            self.suggerimenti_avvocato_nome.pack_forget()
            return
        
        # 🔧 FIX: INVIO chiude il menù e accetta il testo
        if event and event.keysym == 'Return':
            self.suggerimenti_avvocato_nome.pack_forget()
            return
        
        testo = self.entry_avvocato.get()
        if len(testo) < 2:
            self.suggerimenti_avvocato_nome.pack_forget()
            return
        
        avvocati_studio = Config.get_avvocati_studio()
        suggerimenti = []
        testo_lower = testo.lower()
        
        for avvocato in avvocati_studio:
            if testo_lower in avvocato.lower():
                suggerimenti.append(avvocato)
        
        if suggerimenti:
            self.suggerimenti_avvocato_nome.delete(0, tk.END)
            for avvocato in sorted(suggerimenti)[:10]:
                self.suggerimenti_avvocato_nome.insert(tk.END, avvocato)
            self.suggerimenti_avvocato_nome.pack(fill="x", pady=(2, 0))
        else:
            self.suggerimenti_avvocato_nome.pack_forget()
    
    def _on_avvocato_nome_select(self, event=None) -> None:
        """Gestisce la selezione di un avvocato dai suggerimenti."""
        try:
            selection = self.suggerimenti_avvocato_nome.curselection()
            if selection:
                avvocato_selezionato = self.suggerimenti_avvocato_nome.get(selection[0])
                self.entry_avvocato.delete(0, tk.END)
                self.entry_avvocato.insert(0, avvocato_selezionato)
                self.suggerimenti_avvocato_nome.pack_forget()
        except Exception as e:
            print(f"Errore selezione avvocato: {e}")
    
    def _aggiorna_rischio_inerente(self, event=None):
        prestazione = self.combo_prestazione.get()
        prestazioni_veda = Config.get_prestazioni_veda()
        livello = prestazioni_veda.get(prestazione, 2)
        
        livelli_text = {
            1: "NON SIGNIFICATIVO",
            2: "POCO SIGNIFICATIVO",
            3: "ABBASTANZA SIGNIFICATIVO",
            4: "MOLTO SIGNIFICATIVO"
        }
        
        colori = {
            1: "#27AE60",
            2: "#F39C12",
            3: "#E67E22",
            4: "#E74C3C"
        }
        
        self.label_rischio_inerente.config(
            text=f"RISCHIO INERENTE: Livello {livello} - {livelli_text[livello]}",
            fg=colori[livello]
        )
        
        prestazioni_solo_a = Config.get_prestazioni_solo_tabella_a()
        if prestazione in prestazioni_solo_a:
            self.label_solo_tabella_a.config(
                text="⚠️ ATTENZIONE: Prestazione continuativa - Verrà compilata SOLO Tabella A (VEDA pag. 97)"
            )
        else:
            self.label_solo_tabella_a.config(text="")
    
    def _aggiorna_livello_importo(self, event=None):
        self._aggiorna_b3_livello_finale()
        try:
            importo_str = self.entry_importo.get().replace(',', '.').strip()
            if importo_str:
                importo = float(importo_str)
                livello = RiskCalculator.calcola_livello_da_importo(importo)
                livelli_text = {
                    1: "< €50.000 → Livello 1 (NON SIGNIFICATIVO)",
                    2: "€50.000 - €250.000 → Livello 2 (POCO SIGNIFICATIVO)",
                    3: "€250.000 - €1.000.000 → Livello 3 (ABBASTANZA SIGNIFICATIVO)",
                    4: "> €1.000.000 → Livello 4 (MOLTO SIGNIFICATIVO)"
                }
                self.label_importo_livello.config(text=f"🤖 CALCOLO AUTO B3: {livelli_text[livello]} (soglie indicative)")
            else:
                self.label_importo_livello.config(text="")
        except ValueError:
            self.label_importo_livello.config(text="")
    
    def _configure_ttk_style(self) -> None:
        color_scheme = Config.get_color_scheme()
        style = ttk.Style()

        # Configura OptionMenu (dropdown per livelli 1-4) - LEGGIBILE
        style.configure("Custom.TMenubutton",
                       background="#FFFFFF",        # Sfondo bianco
                       foreground="#000000",        # Testo nero
                       borderwidth=2,
                       relief="solid",
                       padding=5)
        style.map("Custom.TMenubutton",
                 background=[("active", color_scheme["select_color"]),
                           ("pressed", color_scheme["select_color"])],
                 foreground=[("active", "#000000")])

        # Configura Combobox per essere leggibile
        style.configure("TCombobox",
                       fieldbackground="#FFFFFF",  # Sfondo bianco campo input
                       background="#FFFFFF",        # Sfondo dropdown
                       foreground="#000000",        # Testo nero
                       arrowcolor="#2C3E50",        # Colore freccia
                       borderwidth=2,
                       relief="solid")
        style.map("TCombobox",
                 fieldbackground=[("readonly", "#FFFFFF")],
                 selectbackground=[("readonly", color_scheme["select_color"])],
                 selectforeground=[("readonly", "#000000")])
    
    def _create_sections(self, parent: tk.Frame) -> None:
        color_scheme = Config.get_color_scheme()
        sections_container = tk.Frame(parent, bg=color_scheme["bg"])
        sections_container.pack(fill="both", expand=True, padx=10, pady=10)
        self._create_section_panel_A(sections_container, "left")
        self._create_section_panel_B(sections_container, "right")
    
    def _create_section_panel_A(self, parent: tk.Frame, side: str) -> None:
        color_scheme = Config.get_color_scheme()
        frame = tk.Frame(parent, bg=color_scheme["frame_bg"], borderwidth=2, relief="groove")
        frame.pack(side=side, fill="both", expand=True, padx=5, pady=5)
        
        tk.Label(frame, text="A. Aspetti connessi al cliente", font=("Helvetica", 12, "bold"), 
                bg=color_scheme["frame_bg"], fg=color_scheme["text_fg"]).pack(pady=10)
        
        self._create_natura_giuridica_section(frame)

        sections_a = Config.get_sections_A_manual()
        for sec in sections_a:
            # Salta A.1 perché è gestito come fattore ibrido automatico/manuale
            if "A.1" in sec["name"]:
                continue
            frame_sec = tk.Frame(frame, bg=color_scheme["frame_bg"], borderwidth=1, relief="solid")
            frame_sec.pack(fill="x", padx=10, pady=5)
            tk.Label(frame_sec, text=sec["name"], bg=color_scheme["frame_bg"],
                    fg=color_scheme["text_fg"], font=("Helvetica", 10, "bold")).pack(anchor="w")
            vars_data = self._create_factors_ui(frame_sec, sec["factors"])
            self.section_vars_A_manual.append({"name": sec["name"], "factors": sec["factors"], **vars_data})
        
        self._create_location_section(frame, "A.4 - Area geografica cliente",
                                     "entry_area_cliente", "label_livello_cliente", "suggerimenti_cliente")
    
    def _create_section_panel_B(self, parent: tk.Frame, side: str) -> None:
        color_scheme = Config.get_color_scheme()
        frame = tk.Frame(parent, bg=color_scheme["frame_bg"], borderwidth=2, relief="groove")
        frame.pack(side=side, fill="both", expand=True, padx=5, pady=5)
        
        tk.Label(frame, text="B. Aspetti connessi all'operazione", font=("Helvetica", 12, "bold"), 
                bg=color_scheme["frame_bg"], fg=color_scheme["text_fg"]).pack(pady=10)
        
        sections_b = Config.get_sections_B_manual()
        for sec in sections_b:
            frame_sec = tk.Frame(frame, bg=color_scheme["frame_bg"], borderwidth=1, relief="solid")
            frame_sec.pack(fill="x", padx=10, pady=5)
            tk.Label(frame_sec, text=sec["name"], bg=color_scheme["frame_bg"],
                    fg=color_scheme["text_fg"], font=("Helvetica", 10, "bold")).pack(anchor="w")
            vars_data = self._create_factors_ui(frame_sec, sec["factors"])
            self.section_vars_B_manual.append({"name": sec["name"], "factors": sec["factors"], **vars_data})

            # Inserisci B.3 dopo B.2
            if sec["name"].startswith("B.2"):
                self._create_b3_ibrido_section(frame)

        self._create_location_section(frame, "B.6 - Area geografica destinazione",
                                     "entry_area_destinazione", "label_livello_destinazione", "suggerimenti_destinazione")
    
    def _create_natura_giuridica_section(self, parent: tk.Frame) -> None:
        color_scheme = Config.get_color_scheme()
        frame_auto = tk.Frame(parent, bg=color_scheme["auto_bg"], borderwidth=2, relief="solid")
        frame_auto.pack(fill="x", padx=10, pady=5)
        
        title_frame = tk.Frame(frame_auto, bg=color_scheme["auto_bg"])
        title_frame.pack(pady=5, padx=10, fill="x")
        
        tk.Label(title_frame, text="A.1 - Natura giuridica", 
                bg=color_scheme["auto_bg"], fg=color_scheme["auto_label"], 
                font=("Helvetica", 10, "bold")).pack(side="left")
        
        info_label = tk.Label(title_frame, text=" ℹ️", bg=color_scheme["auto_bg"], 
                             font=("Helvetica", 12), cursor="question_arrow")
        info_label.pack(side="left", padx=5)
        
        tooltip_text = """RICERCA INTELLIGENTE NATURA GIURIDICA:

🏦 DATABASE STANDARD:
- 20+ Banche italiane (Intesa, UniCredit, MPS...)
- 10+ Assicurazioni (Generali, Allianz, UnipolSai...)
- 30+ Grandi aziende quotate (ENI, ENEL, FIAT...)
- Tutte le PA (Comuni, Regioni, Ministeri...)

🏢 DATABASE CLIENTI STUDIO:
- Lista completa clienti dello studio
- SENZA livello pre-assegnato
- Richiede selezione manuale

🔍 FUNZIONAMENTO:
1. Database standard → Livello automatico ✅
2. Clienti studio → Seleziona livello manualmente ⚠️
3. Keywords (SRL, Trust...) → Suggerimento categoria

ESEMPI:
- "Intesa Sanpaolo" → Livello 1 (automatico)
- "Garage Raw" → Cliente studio, scegli livello
- "Mario Rossi SRL" → Livello 3 (da keyword)"""
        
        ToolTip(info_label, tooltip_text, delay=300)
        
        input_frame = tk.Frame(frame_auto, bg=color_scheme["auto_bg"])
        input_frame.pack(padx=10, pady=5, fill="x")
        
        tk.Label(input_frame, text="Nome Cliente/Ragione Sociale:", 
                bg=color_scheme["auto_bg"], fg=color_scheme["text_fg"], 
                font=("Helvetica", 9)).pack(side="left", padx=(0, 5))
        
        entry = tk.Entry(input_frame, bg=color_scheme["input_bg"], fg="#000000", 
                        borderwidth=2, width=40, font=("Helvetica", 10))
        entry.pack(side="left", padx=5)
        self.entry_natura_giuridica = entry
        
        search_icon = tk.Label(input_frame, text="🔍", bg=color_scheme["auto_bg"], 
                              font=("Helvetica", 14))
        search_icon.pack(side="left", padx=2)
        
        listbox = tk.Listbox(frame_auto, height=6, bg="#FFFFFF", fg="#000000", 
                            font=("Helvetica", 9), selectbackground=color_scheme["select_color"])
        self.suggerimenti_natura = listbox
        
        label_livello = tk.Label(frame_auto, text="", bg=color_scheme["auto_bg"], 
                                fg=color_scheme["auto_label"], 
                                font=("Helvetica", 10, "bold"), pady=5)
        label_livello.pack()
        self.label_natura_livello = label_livello
        
        separator = tk.Frame(frame_auto, height=2, bg="#BDC3C7")
        separator.pack(fill="x", padx=20, pady=10)
        
        fallback_frame = tk.Frame(frame_auto, bg=color_scheme["auto_bg"])
        fallback_frame.pack(padx=10, pady=5, fill="x")
        
        tk.Label(fallback_frame, text="⚙️ Seleziona categoria manuale:", 
                bg=color_scheme["auto_bg"], fg=color_scheme["text_fg"], 
                font=("Helvetica", 9, "italic")).pack(side="left", padx=(0, 10))
        
        categorie_manuali = Config.get_natura_giuridica_categorie_manuali()
        # Filtra chiavi che iniziano con "_" (sono metadati, non categorie)
        categorie_valide = [k for k in categorie_manuali.keys() if not k.startswith("_")]
        combo_var = tk.StringVar(value="")
        combo_fallback = ttk.Combobox(fallback_frame, textvariable=combo_var,
                                     values=categorie_valide,
                                     state="readonly", width=65, font=("Helvetica", 9))
        combo_fallback.pack(side="left")
        self.combo_natura_fallback = combo_var
        
        combo_fallback.bind('<<ComboboxSelected>>', lambda e: self._on_natura_fallback_select())
        entry.bind('<KeyRelease>', lambda e: self._on_natura_keyrelease())
        listbox.bind('<<ListboxSelect>>', lambda e: self._on_natura_select())
        
        # 🔧 FIX: Chiude menù con ESC, INVIO o click fuori
        entry.bind('<Escape>', lambda e: self.suggerimenti_natura.pack_forget())
        entry.bind('<Return>', lambda e: self.suggerimenti_natura.pack_forget())
        entry.bind('<FocusOut>', lambda e: self.suggerimenti_natura.pack_forget())
        
        help_label = tk.Label(frame_auto,
                             text="Digita per cercare: Database standard",
                             bg=color_scheme["auto_bg"], fg="#7F8C8D",
                             font=("Helvetica", 8, "italic"))
        help_label.pack(pady=(0, 5))

        # === MODIFICA 1 VEDA: Fattori soggettivi A.1 ===
        separator2 = tk.Frame(frame_auto, height=2, bg="#BDC3C7")
        separator2.pack(fill="x", padx=20, pady=10)

        tk.Label(frame_auto,
                text="Fattori soggettivi aggiuntivi (valutazione dell'avvocato):",
                bg=color_scheme["auto_bg"], fg=color_scheme["text_fg"],
                font=("Helvetica", 10, "bold")).pack(anchor="w", padx=10, pady=(5, 2))

        # Carica i fattori A.1 dalla configurazione
        fattori_a1 = Config.get_sections_A_manual()
        a1_factors = None
        for sec in fattori_a1:
            if sec["name"].startswith("A.1"):
                a1_factors = sec["factors"]
                break

        if a1_factors:
            factors_frame = tk.Frame(frame_auto, bg=color_scheme["auto_bg"])
            factors_frame.pack(padx=10, fill="x", pady=5)
            vars_data = self._create_factors_ui(factors_frame, a1_factors)
            self.a1_manual_factors = {"name": "A.1 - Natura giuridica",
                                     "factors": a1_factors, **vars_data}

    def _auto_detect_natura_giuridica(self, testo: str) -> None:
        """Rileva automaticamente il livello di natura giuridica dal testo."""
        import re
        color_scheme = Config.get_color_scheme()
        database = Config.get_natura_giuridica_database()

        # Normalizza il testo: rimuove punti, virgole, trattini
        def normalizza(s: str) -> str:
            # Rimuove punti, virgole, trattini (ma mantiene gli spazi tra parole)
            s = re.sub(r'[.,\-]', '', s.lower().strip())
            return s

        # Trova il testo del dropdown per un dato livello
        def get_dropdown_text_for_level(livello: int) -> str:
            categorie_manuali = Config.get_natura_giuridica_categorie_manuali()
            for key, val in categorie_manuali.items():
                if val == livello:
                    return key
            return ""

        testo_norm = normalizza(testo)

        # Cerca match esatto nel database
        for nome, (livello, descrizione) in database.items():
            nome_norm = normalizza(nome)
            if nome_norm in testo_norm:
                livello_text = {1: "NON SIGNIFICATIVO", 2: "POCO SIGNIFICATIVO",
                              3: "ABBASTANZA SIGNIFICATIVO", 4: "MOLTO SIGNIFICATIVO"}
                self.label_natura_livello.config(
                    text=f"🤖 RILEVATO AUTOMATICAMENTE - NON SELEZIONARE NULLA DAL MENÙ A TENDINA SOTTOSTANTE: Livello {livello} - {livello_text[livello]}",
                    fg=color_scheme["auto_label"]
                )
                # Seleziona automaticamente nel dropdown
                dropdown_text = get_dropdown_text_for_level(livello)
                self.combo_natura_fallback.set(dropdown_text)
                return

        # Se non trovato nel database, cerca per keyword
        keywords = Config.get_natura_giuridica_keywords()
        for livello, keyword_list in keywords.items():
            for keyword in keyword_list:
                keyword_norm = normalizza(keyword)
                if keyword_norm in testo_norm:
                    livello_text = {1: "NON SIGNIFICATIVO", 2: "POCO SIGNIFICATIVO",
                                  3: "ABBASTANZA SIGNIFICATIVO", 4: "MOLTO SIGNIFICATIVO"}
                    self.label_natura_livello.config(
                        text=f"🤖 RILEVATO AUTOMATICAMENTE: Livello {livello} - {livello_text[livello]}",
                        fg=color_scheme["auto_label"]
                    )
                    # Seleziona automaticamente nel dropdown
                    dropdown_text = get_dropdown_text_for_level(livello)
                    self.combo_natura_fallback.set(dropdown_text)
                    return

        # Se nessun match, resetta il label e il dropdown
        self.label_natura_livello.config(text="⚠️ Nessun match automatico - Seleziona manualmente", fg="#E74C3C")
        self.combo_natura_fallback.set("")

    def _on_natura_fallback_select(self) -> None:
        selected = self.combo_natura_fallback.get()
        if selected:
            categorie_manuali = Config.get_natura_giuridica_categorie_manuali()
            livello = categorie_manuali[selected]
            livello_text = {1: "NON SIGNIFICATIVO", 2: "POCO SIGNIFICATIVO",
                          3: "ABBASTANZA SIGNIFICATIVO", 4: "MOLTO SIGNIFICATIVO"}
            self.label_natura_livello.config(
                text=f"SELEZIONE MANUALE: Livello {livello} - {livello_text[livello]}",
                fg="#27AE60")
    
    def _on_natura_keyrelease(self) -> None:
        testo = self.entry_natura_giuridica.get()
        if len(testo) < 2:
            self.suggerimenti_natura.pack_forget()
            self.label_natura_livello.config(text="")
            return

        # Auto-rileva il livello mentre l'utente digita
        self._auto_detect_natura_giuridica(testo)

        suggerimenti = RiskCalculator.cerca_natura_giuridica(testo)
        if suggerimenti:
            self.suggerimenti_natura.delete(0, tk.END)
            for nome, livello, descrizione, is_cliente_studio in suggerimenti:
                if is_cliente_studio:
                    self.suggerimenti_natura.insert(tk.END, 
                        f"{nome} [Cliente studio - SELEZIONA LIVELLO ⬇️]")
                else:
                    livello_text = {1: "NON SIG", 2: "POCO SIG", 3: "ABBAST SIG", 4: "MOLTO SIG"}
                    self.suggerimenti_natura.insert(tk.END, 
                        f"{nome} [Liv {livello} - {livello_text[livello]}] - {descrizione}")
            self.suggerimenti_natura.pack(padx=10, pady=(0, 5), fill="x")
        else:
            self.suggerimenti_natura.pack_forget()
            self.label_natura_livello.config(
                text="⚠️ Non trovato. Usa il menu categoria manuale qui sotto ⬇️")
    
    def _on_natura_select(self) -> None:
        try:
            selection = self.suggerimenti_natura.curselection()
            if selection:
                value = self.suggerimenti_natura.get(selection[0])
                color_scheme = Config.get_color_scheme()
                
                if "[Cliente studio - SELEZIONA LIVELLO" in value:
                    nome = value.split(" [Cliente studio")[0]
                    self.entry_natura_giuridica.delete(0, tk.END)
                    self.entry_natura_giuridica.insert(0, nome)
                    self.suggerimenti_natura.pack_forget()
                    self.label_natura_livello.config(
                        text="⚠️ CLIENTE STUDIO: Seleziona livello nel menu qui sotto ⬇️",
                        fg="#E74C3C"
                    )
                    self.combo_natura_fallback.set("")
                else:
                    nome = value.split(" [Liv ")[0]
                    livello = int(value.split("[Liv ")[1].split(" -")[0])
                    self.entry_natura_giuridica.delete(0, tk.END)
                    self.entry_natura_giuridica.insert(0, nome)
                    livello_text = {1: "NON SIGNIFICATIVO", 2: "POCO SIGNIFICATIVO", 
                                  3: "ABBASTANZA SIGNIFICATIVO", 4: "MOLTO SIGNIFICATIVO"}
                    self.label_natura_livello.config(
                        text=f"🤖 RILEVATO AUTOMATICAMENTE: Livello {livello} - {livello_text[livello]}",
                        fg=color_scheme["auto_label"]
                    )
                    self.combo_natura_fallback.set("")
                    self.suggerimenti_natura.pack_forget()
        except Exception as e:
            print(f"Errore selezione natura: {e}")
    
    def _create_b3_ibrido_section(self, parent: tk.Frame) -> None:
        color_scheme = Config.get_color_scheme()
        frame_b3 = tk.Frame(parent, bg=color_scheme["auto_bg"], borderwidth=2, relief="solid")
        frame_b3.pack(fill="x", padx=10, pady=5)
        
        title_frame = tk.Frame(frame_b3, bg=color_scheme["auto_bg"])
        title_frame.pack(pady=5, padx=10, fill="x")
        
        tk.Label(title_frame, text="B.3 - Ammontare operazione [IBRIDO - AUTO + VALUTAZIONE]", 
                bg=color_scheme["auto_bg"], fg=color_scheme["auto_label"], 
                font=("Helvetica", 10, "bold")).pack(side="left")
        
        info_label = tk.Label(title_frame, text=" ℹ️", bg=color_scheme["auto_bg"], 
                             font=("Helvetica", 12), cursor="question_arrow")
        info_label.pack(side="left", padx=5)
        
        tooltip_text = """CALCOLO IBRIDO B.3 - AMMONTARE:

1️⃣ LIVELLO BASE (automatico da importo - soglie indicative):
- < €50.000 → Livello 1
- €50.000 - €250.000 → Livello 2
- €250.000 - €1.000.000 → Livello 3
- > €1.000.000 → Livello 4

⚠️ ATTENZIONE: Soglie NON presenti in VEDA ufficiale
(basate su prassi comune di mercato)

2️⃣ FATTORI AGGRAVANTI (valutazione professionale):
- Incongruenza con profilo cliente → +1 livello
- Frazionamenti artificiosi → +1 livello
- Altro fattore anomalo → +1 livello

3️⃣ LIVELLO FINALE = BASE + AGGRAVANTI
(massimo livello 4)

ESEMPIO:
Importo €80.000 → Base livello 2
+ Incongruo con profilo → +1
= LIVELLO FINALE 3"""
        
        ToolTip(info_label, tooltip_text, delay=300)
        
        self.label_b3_base = tk.Label(frame_b3, text="", bg=color_scheme["auto_bg"], 
                                     fg="#3498DB", font=("Helvetica", 9, "bold"))
        self.label_b3_base.pack(pady=5)
        
        separator = tk.Frame(frame_b3, height=1, bg="#BDC3C7")
        separator.pack(fill="x", padx=20, pady=5)
        
        tk.Label(frame_b3, text="Verifica fattori aggravanti:", 
                bg=color_scheme["auto_bg"], fg=color_scheme["text_fg"], 
                font=("Helvetica", 9, "bold")).pack(anchor="w", padx=10, pady=(5, 2))
        
        self.check_b3_incongruo = tk.BooleanVar()
        check1 = tk.Checkbutton(frame_b3, 
                               text="Importo incongruo rispetto al profilo economico-finanziario del cliente",
                               variable=self.check_b3_incongruo, bg=color_scheme["auto_bg"],
                               fg=color_scheme["text_fg"], selectcolor=color_scheme["select_color"],
                               command=self._aggiorna_b3_livello_finale, anchor="w", wraplength=450)
        check1.pack(fill="x", padx=20, pady=2)
        
        self.check_b3_frazionamenti = tk.BooleanVar()
        check2 = tk.Checkbutton(frame_b3, text="Presenza di frazionamenti artificiosi",
                               variable=self.check_b3_frazionamenti, bg=color_scheme["auto_bg"],
                               fg=color_scheme["text_fg"], selectcolor=color_scheme["select_color"],
                               command=self._aggiorna_b3_livello_finale, anchor="w")
        check2.pack(fill="x", padx=20, pady=2)
        
        self.check_b3_altro = tk.BooleanVar()
        check3 = tk.Checkbutton(frame_b3, text="Altro fattore anomalo",
                               variable=self.check_b3_altro, bg=color_scheme["auto_bg"],
                               fg=color_scheme["text_fg"], selectcolor=color_scheme["select_color"],
                               command=self._aggiorna_b3_livello_finale, anchor="w")
        check3.pack(fill="x", padx=20, pady=2)
        
        separator2 = tk.Frame(frame_b3, height=1, bg="#BDC3C7")
        separator2.pack(fill="x", padx=20, pady=5)
        
        self.label_b3_livello_finale = tk.Label(frame_b3, text="", bg=color_scheme["auto_bg"], 
                                               fg=color_scheme["auto_label"], 
                                               font=("Helvetica", 10, "bold"), pady=5)
        self.label_b3_livello_finale.pack()
        
        help_label = tk.Label(frame_b3, 
                             text="Il livello finale è calcolato automaticamente: Base + Aggravanti (max 4)", 
                             bg=color_scheme["auto_bg"], fg="#7F8C8D",
                             font=("Helvetica", 8, "italic"))
        help_label.pack(pady=(0, 5))
    
    def _aggiorna_b3_livello_finale(self):
        try:
            importo_str = self.entry_importo.get().replace(',', '.').strip()
            if not importo_str:
                self.label_b3_base.config(text="⚠️ Inserisci l'importo per calcolare il livello base")
                self.label_b3_livello_finale.config(text="")
                return
            importo = float(importo_str)
            livello_base = RiskCalculator.calcola_livello_da_importo(importo)
            fascia_text = {
                1: "< €50.000",
                2: "€50.000 - €250.000",
                3: "€250.000 - €1.000.000",
                4: "> €1.000.000"
            }
            self.label_b3_base.config(text=f"Livello BASE da importo: {livello_base} ({fascia_text[livello_base]}) [soglie indicative]")
            aggravanti = 0
            if self.check_b3_incongruo.get():
                aggravanti += 1
            if self.check_b3_frazionamenti.get():
                aggravanti += 1
            if self.check_b3_altro.get():
                aggravanti += 1
            livello_finale = min(livello_base + aggravanti, 4)
            livelli_text = {
                1: "NON SIGNIFICATIVO",
                2: "POCO SIGNIFICATIVO",
                3: "ABBASTANZA SIGNIFICATIVO",
                4: "MOLTO SIGNIFICATIVO"
            }
            colori = {
                1: "#27AE60",
                2: "#F39C12",
                3: "#E67E22",
                4: "#E74C3C"
            }
            if aggravanti > 0:
                self.label_b3_livello_finale.config(
                    text=f"LIVELLO FINALE B.3: {livello_finale} - {livelli_text[livello_finale]} (Base {livello_base} + {aggravanti} aggravanti)",
                    fg=colori[livello_finale]
                )
            else:
                self.label_b3_livello_finale.config(
                    text=f"LIVELLO FINALE B.3: {livello_finale} - {livelli_text[livello_finale]}",
                    fg=colori[livello_finale]
                )
        except ValueError:
            self.label_b3_base.config(text="Importo non valido")
            self.label_b3_livello_finale.config(text="")
    
    def _create_location_section(self, parent: tk.Frame, title: str, 
                                 entry_attr: str, label_attr: str, listbox_attr: str) -> None:
        color_scheme = Config.get_color_scheme()
        frame_auto = tk.Frame(parent, bg=color_scheme["auto_bg"], borderwidth=2, relief="solid")
        frame_auto.pack(fill="x", padx=10, pady=5)
        
        title_frame = tk.Frame(frame_auto, bg=color_scheme["auto_bg"])
        title_frame.pack(pady=5, padx=10, fill="x")
        
        tk.Label(title_frame, text=title, bg=color_scheme["auto_bg"], 
                fg=color_scheme["auto_label"], font=("Helvetica", 10, "bold")).pack(side="left")
        
        info_label = tk.Label(title_frame, text=" ℹ️", bg=color_scheme["auto_bg"], 
                             font=("Helvetica", 12), cursor="question_arrow")
        info_label.pack(side="left", padx=5)
        
        tooltip_text = """RICERCA INTELLIGENTE LUOGHI:

Inizia a digitare il nome della provincia o del paese.
Il sistema cercherà automaticamente nel database di:

🇮🇹 108 province italiane 
   (rischio uso contante - Fonte: Banca d'Italia/MEF 2019)
   Riferimento: fisco7.it/mappatura-rischio

🌍 50+ paesi internazionali 
   (rischio AML - Fonte: Linee Guida VEDA + Reg. UE 2019/1326)

Il livello di rischio verrà assegnato automaticamente.

ESEMPI:
- "Milano" → Livello 4 (rischio alto contante)
- "Afghanistan" → Livello 4 (paese alto rischio UE)
- "Genova" → Livello 1 (rischio basso)
- "USA" → Livello 1 (basso rischio)

CHIUSURA MENÙ: INVIO, ESC o click fuori"""
        
        ToolTip(info_label, tooltip_text, delay=300)
        
        input_frame = tk.Frame(frame_auto, bg=color_scheme["auto_bg"])
        input_frame.pack(padx=10, pady=5, fill="x")
        
        tk.Label(input_frame, text="Provincia/Città/Paese:", bg=color_scheme["auto_bg"], 
                fg=color_scheme["text_fg"], font=("Helvetica", 9)).pack(side="left", padx=(0, 5))
        
        entry = tk.Entry(input_frame, bg=color_scheme["input_bg"], fg="#000000", 
                        borderwidth=2, width=40, font=("Helvetica", 10))
        entry.pack(side="left", padx=5)
        setattr(self, entry_attr, entry)
        
        search_icon = tk.Label(input_frame, text="🔍", bg=color_scheme["auto_bg"], 
                              font=("Helvetica", 14))
        search_icon.pack(side="left", padx=2)
        
        listbox = tk.Listbox(frame_auto, height=6, bg="#FFFFFF", fg="#000000", 
                            font=("Helvetica", 9), selectbackground=color_scheme["select_color"])
        setattr(self, listbox_attr, listbox)
        
        label_livello = tk.Label(frame_auto, text="", bg=color_scheme["auto_bg"], 
                                fg=color_scheme["auto_label"], 
                                font=("Helvetica", 10, "bold"), pady=5)
        label_livello.pack()
        setattr(self, label_attr, label_livello)
        
        entry.bind('<KeyRelease>', lambda e: self._on_location_keyrelease(e, entry, listbox, label_livello))
        listbox.bind('<<ListboxSelect>>', lambda e: self._on_location_select(listbox, entry, label_livello))
        # 🔧 FIX: Chiude menù quando si clicca fuori dal campo
        entry.bind('<FocusOut>', lambda e: listbox.pack_forget())
        
        help_label = tk.Label(frame_auto, 
                             text="Inizia a digitare per vedere i suggerimenti automatici (Province ITA: Banca d'Italia / Paesi: VEDA) | INVIO/ESC per chiudere", 
                             bg=color_scheme["auto_bg"], fg="#7F8C8D",
                             font=("Helvetica", 8, "italic"))
        help_label.pack(pady=(0, 5))
    
    def _on_location_keyrelease(self, event, entry: tk.Entry, listbox: tk.Listbox, label: tk.Label) -> None:
        # 🔧 FIX: ESC chiude il menù
        if event and event.keysym == 'Escape':
            listbox.pack_forget()
            return

        # 🔧 FIX: INVIO chiude il menù e accetta il testo
        if event and event.keysym == 'Return':
            listbox.pack_forget()
            return
        
        testo = entry.get()
        if len(testo) < 2:
            listbox.pack_forget()
            label.config(text="")
            return
        suggerimenti = RiskCalculator.cerca_luogo(testo)
        if suggerimenti:
            listbox.delete(0, tk.END)
            for luogo, livello in suggerimenti:
                livello_text = {1: "BASSO", 2: "MEDIO", 3: "MEDIO-ALTO", 4: "ALTO"}
                listbox.insert(tk.END, f"{luogo} [Livello {livello} - {livello_text[livello]}]")
            listbox.pack(padx=10, pady=(0, 5), fill="x")
        else:
            listbox.pack_forget()
            label.config(text="⚠️ Luogo non trovato nel database. Inserisci il livello manualmente (1-4)")
    
    def _on_location_select(self, listbox: tk.Listbox, entry: tk.Entry, label: tk.Label) -> None:
        try:
            selection = listbox.curselection()
            if selection:
                value = listbox.get(selection[0])
                luogo = value.split(" [")[0]
                entry.delete(0, tk.END)
                entry.insert(0, luogo)
                livello = RiskCalculator.get_livello_paese(luogo)
                livello_text = {1: "NON SIGNIFICATIVO", 2: "POCO SIGNIFICATIVO", 
                              3: "ABBASTANZA SIGNIFICATIVO", 4: "MOLTO SIGNIFICATIVO"}
                label.config(text=f"🤖 RILEVATO AUTOMATICAMENTE: Livello {livello} - {livello_text[livello]}")
                listbox.pack_forget()
        except Exception as e:
            print(f"Errore selezione: {e}")
    
    def _create_factors_ui(self, parent: tk.Frame, factors: List[str]) -> Dict:
        color_scheme = Config.get_color_scheme()
        vars_list = []
        level_vars_list = []
        for factor in factors:
            container = tk.Frame(parent, bg=color_scheme["frame_bg"])
            container.pack(fill="x", pady=2, padx=5)
            var = tk.BooleanVar()
            level_var = tk.IntVar(value=1)
            max_length = 100
            display_text = factor[:max_length] + "..." if len(factor) > max_length else factor
            check = tk.Checkbutton(container, text=display_text, variable=var, 
                                  bg=color_scheme["frame_bg"], fg=color_scheme["text_fg"], 
                                  selectcolor=color_scheme["select_color"], wraplength=500, 
                                  justify="left", anchor="w")
            check.pack(side="left", padx=5, fill="x", expand=True)
            if len(factor) > max_length:
                ToolTip(check, factor, delay=500)
            level_menu = ttk.OptionMenu(container, level_var, 1, 1, 2, 3, 4, style="Custom.TMenubutton")
            level_menu.pack_forget()
            check.config(command=lambda v=var, m=level_menu: self._toggle_level(v, m))
            vars_list.append(var)
            level_vars_list.append(level_var)
        return {"vars": vars_list, "level_vars": level_vars_list}
    
    @staticmethod
    def _toggle_level(var: tk.BooleanVar, level_menu: ttk.OptionMenu) -> None:
        if var.get():
            level_menu.pack(side="left", padx=5)
        else:
            level_menu.pack_forget()
    
    def _create_action_buttons(self, parent: tk.Frame) -> None:
        color_scheme = Config.get_color_scheme()
        btn_frame = tk.Frame(parent, bg=color_scheme["bg"])
        btn_frame.pack(fill="x", pady=10)
        buttons = [("Valuta Rischio", self.valuta_rischio), ("Esporta Word", self.esporta_word)]
        for text, command in buttons:
            btn = tk.Button(btn_frame, text=text, command=command, bg=color_scheme["button_bg"], 
                          fg=color_scheme["button_fg"], font=("Helvetica", 11, "bold"), 
                          padx=20, pady=8, relief="raised", borderwidth=2, cursor="hand2", 
                          activebackground=color_scheme["button_active"], activeforeground="white")
            btn.pack(side="left", padx=5)
            
            def on_enter(e, button=btn):
                button['background'] = color_scheme["button_hover"]
            
            def on_leave(e, button=btn):
                button['background'] = color_scheme["button_bg"]
            
            btn.bind("<Enter>", on_enter)
            btn.bind("<Leave>", on_leave)
    
    def _get_automatic_sections_A(self) -> List[Dict]:
        sections = []
        natura_text = self.entry_natura_giuridica.get().strip()
        natura_level = 0
        descrizione = ""
        database = Config.get_natura_giuridica_database()
        for nome, (lvl, desc) in database.items():
            if nome.lower() in natura_text.lower():
                natura_level = lvl
                descrizione = desc
                break
        if natura_level == 0:
            natura_lower = natura_text.lower()
            keywords = Config.get_natura_giuridica_keywords()
            for lvl, keyword_list in keywords.items():
                for keyword in keyword_list:
                    if keyword in natura_lower:
                        natura_level = lvl
                        break
                if natura_level > 0:
                    break
        if natura_level == 0 and self.combo_natura_fallback and hasattr(self.combo_natura_fallback, 'get'):
            selected_cat = self.combo_natura_fallback.get()
            if selected_cat:
                categorie_manuali = Config.get_natura_giuridica_categorie_manuali()
                natura_level = categorie_manuali.get(selected_cat, 3)
                descrizione = ""
        if natura_level == 0:
            natura_level = 3
            descrizione = "Livello default - verificare manualmente"

        # A.1 usa solo il livello della natura giuridica selezionata, SENZA medie
        value_text = f"{natura_text} ({descrizione})" if descrizione else natura_text

        sections.append({
            "name": "A.1 - Natura giuridica",
            "value": value_text,
            "level": natura_level
        })
        area_cliente = self.entry_area_cliente.get().strip()
        if area_cliente:
            area_level = RiskCalculator.get_livello_paese(area_cliente)
            if area_level == 0:
                area_level = 1
            sections.append({
                "name": "A.4 - Area geografica cliente",
                "value": area_cliente,
                "level": area_level
            })
        else:
            sections.append({
                "name": "A.4 - Area geografica cliente",
                "value": "Non specificato",
                "level": 1
            })
        return sections
    
    def _get_automatic_sections_B(self, importo: Optional[float]) -> List[Dict]:
        sections = []
        if importo is not None:
            importo_level_base = RiskCalculator.calcola_livello_da_importo(importo)

            # B.3 usa solo il livello base dell'importo, SENZA medie
            # I flag Incongruo/Frazionamenti/Altro sono solo informativi

            fascia_text = {
                1: "< €50.000",
                2: "€50.000 - €250.000",
                3: "€250.000 - €1.000.000",
                4: "> €1.000.000"
            }

            descrizione = f"€ {importo:,.2f} ({fascia_text[importo_level_base]})"

            sections.append({
                "name": "B.3 - Ammontare operazione",
                "value": descrizione,
                "level": importo_level_base
            })
        else:
            sections.append({
                "name": "B.3 - Ammontare operazione",
                "value": "Non specificato",
                "level": 0
            })
        area_dest = self.entry_area_destinazione.get().strip()
        if area_dest:
            dest_level = RiskCalculator.get_livello_paese(area_dest)
            if dest_level == 0:
                dest_level = 1
            sections.append({
                "name": "B.6 - Area geografica destinazione",
                "value": area_dest,
                "level": dest_level
            })
        else:
            sections.append({
                "name": "B.6 - Area geografica destinazione",
                "value": "Non specificato",
                "level": 1
            })
        return sections
    
    def valuta_rischio(self) -> None:
        try:
            data_valutazione = self.entry_data.get().strip()
            if not data_valutazione:
                messagebox.showwarning("Attenzione", "Inserire la data di valutazione!")
                return
            
            cliente = self.entry_cliente.get().strip()
            titolare_effettivo = self.entry_titolare.get().strip()
            scopo_operazione = self.text_scopo.get("1.0", tk.END).strip()
            descrizione_attivita = self.text_attivita.get("1.0", tk.END).strip()
            avvocato = self.entry_avvocato.get().strip()
            
            if not cliente:
                messagebox.showwarning("Attenzione", "Inserire il nome del cliente!")
                return
            
            prestazione = self.combo_prestazione.get()
            prestazioni_veda = Config.get_prestazioni_veda()
            rischio_inerente = prestazioni_veda.get(prestazione, 2)
            prestazioni_solo_a = Config.get_prestazioni_solo_tabella_a()
            usa_solo_tabella_a = prestazione in prestazioni_solo_a
            
            importo_str = self.entry_importo.get().replace(',', '.').strip()
            importo = float(importo_str) if importo_str else None
            anomalie = RiskCalculator.valida_anomalia_importo(importo)
            
            sections_A_auto = self._get_automatic_sections_A()
            sections_B_auto = self._get_automatic_sections_B(importo) if not usa_solo_tabella_a else []
            
            total_A, sub_medias_A = RiskCalculator.calcola_totale_sezioni(
                self.section_vars_A_manual, sections_A_auto
            )
            
            if usa_solo_tabella_a:
                total_B = 0.0
                sub_medias_B = []
                num_fattori_a = len(sections_A_auto) + len(self.section_vars_A_manual)
                rischio_specifico = round(total_A / num_fattori_a, 2)
                nota_calcolo = " (SOLO TABELLA A - prestazione continuativa VEDA pag. 97)"
            else:
                total_B, sub_medias_B = RiskCalculator.calcola_totale_sezioni(
                    self.section_vars_B_manual, sections_B_auto
                )
                rischio_specifico = round((total_A + total_B) / 10, 2)
                num_fattori_a = 0
                nota_calcolo = ""
            
            inerente_ponderato = round(rischio_inerente * 0.3, 2)
            specifico_ponderato = round(rischio_specifico * 0.7, 2)
            somma = round(inerente_ponderato + specifico_ponderato, 2)
            livello = RiskCalculator.calcola_livello_rischio(somma)
            
            importo_display = f"€ {importo:,.2f}" if importo else "Non specificato"
            
            risultato = f"""═══════════════════════════════════════════════════
  RISULTATO PROFILATURA RISCHIO AML
  Data: {data_valutazione}
═══════════════════════════════════════════════════

Cliente: {cliente}
Titolare Effettivo: {titolare_effettivo}
Scopo Operazione: {scopo_operazione}
Importo: {importo_display}
Prestazione: {prestazione}"""

            if usa_solo_tabella_a:
                risultato += f"""
Prestazione continuativa - Compilata SOLO Tabella A
(La Tabella B non è applicabile per questa prestazione)"""

            risultato += f"""
            """
            
            for sec in sections_A_auto:
                risultato += f"\n{sec['name']}: {sec['value']}"
            for sec in sections_B_auto:
                risultato += f"\n{sec['name']}: {sec['value']} [Livello {sec['level']}]"
            
            risultato += f"""

─────────────────────────────────────────────────
CALCOLI:
─────────────────────────────────────────────────
Totale A (Cliente): {total_A:.2f}"""

            if not usa_solo_tabella_a:
                risultato += f"""
Totale B (Operazione): {total_B:.2f}
Rischio Specifico: {rischio_specifico:.2f} = (A + B) / 10"""
            else:
                risultato += f"""
Totale B (Operazione): NON COMPILATA
Rischio Specifico: {rischio_specifico:.2f} = A / {num_fattori_a} fattori{nota_calcolo}"""

            risultato += f"""
Rischio Inerente: {rischio_inerente}

Ponderazione:
  Inerente × 0.3 = {inerente_ponderato:.2f}
  Specifico × 0.7 = {specifico_ponderato:.2f}

═══════════════════════════════════════════════════
  SOMMA PONDERATA: {somma:.2f}
  LIVELLO RISCHIO: {livello}
═══════════════════════════════════════════════════

{anomalie}

Avv. {avvocato}"""
            
            self.label_risultato.config(text=risultato)

            # Gestione Rischio Rafforzato (RE > 3.5)
            dichiarazione_rafforzata = None
            if somma > 3.5:
                dichiarazione_rafforzata = self._mostra_dialog_rischio_rafforzato(somma, livello)
                if dichiarazione_rafforzata is None:
                    # L'utente ha chiuso il dialog senza accettare
                    return

            self.dati_export = {
                "data_valutazione": data_valutazione,
                "cliente": cliente,
                "titolare_effettivo": titolare_effettivo,
                "scopo_operazione": scopo_operazione,
                "descrizione_attivita": descrizione_attivita,
                "rischio_inerente": rischio_inerente,
                "prestazione_veda": prestazione,
                "avvocato": avvocato,
                "importo": importo,
                "total_A": total_A,
                "total_B": total_B,
                "rischio_specifico": rischio_specifico,
                "inerente_ponderato": inerente_ponderato,
                "specifico_ponderato": specifico_ponderato,
                "somma": somma,
                "livello": livello,
                "anomalie": anomalie,
                "sections_A_manual": self.section_vars_A_manual,
                "sections_A_auto": sections_A_auto,
                "sections_B_manual": self.section_vars_B_manual,
                "sections_B_auto": sections_B_auto,
                "sub_medias_A": sub_medias_A,
                "sub_medias_B": sub_medias_B,
                "usa_solo_tabella_a": usa_solo_tabella_a,
                "num_fattori_a": num_fattori_a if usa_solo_tabella_a else 0,
                "dichiarazione_rafforzata": dichiarazione_rafforzata
            }
            
            messagebox.showinfo("Successo", "Valutazione completata con successo!")
        
        except ValueError as e:
            messagebox.showerror("Errore", f"Controlla gli input:\n- Livelli devono essere tra 1 e 4\n- Importo deve essere un numero valido\n\nDettaglio: {str(e)}")
        except Exception as e:
            messagebox.showerror("Errore", f"Errore imprevisto: {str(e)}")

    def _mostra_dialog_rischio_rafforzato(self, somma_ponderata: float, livello: str) -> Optional[Dict]:
        """
        Mostra un dialog modale per Rischio Rafforzato (RE > 3.5).
        Richiede l'accettazione esplicita degli obblighi di Adeguata Verifica Rafforzata.
        Restituisce un dizionario con timestamp e dichiarazione, oppure None se rifiutato.
        """
        dialog = tk.Toplevel(self.root)
        dialog.title("⚠️ RISCHIO RAFFORZATO - Adeguata Verifica Rafforzata Obbligatoria")
        dialog.geometry("800x600")
        dialog.resizable(False, False)
        dialog.grab_set()  # Modale
        dialog.transient(self.root)

        # Centra il dialog
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() // 2) - (800 // 2)
        y = (dialog.winfo_screenheight() // 2) - (600 // 2)
        dialog.geometry(f"+{x}+{y}")

        color_scheme = Config.get_color_scheme()
        dialog.configure(bg="#FFF3CD")  # Giallo warning

        # Header
        header_frame = tk.Frame(dialog, bg="#FFC107", pady=15)
        header_frame.pack(fill="x")

        tk.Label(header_frame, text="⚠️ ATTENZIONE: RISCHIO RAFFORZATO RILEVATO",
                font=("Helvetica", 16, "bold"), bg="#FFC107", fg="#000000").pack()

        tk.Label(header_frame, text=f"Rischio Effettivo Calcolato: {somma_ponderata:.2f} - {livello}",
                font=("Helvetica", 12), bg="#FFC107", fg="#000000").pack(pady=5)

        # Body
        body_frame = tk.Frame(dialog, bg="#FFF3CD", padx=30, pady=20)
        body_frame.pack(fill="both", expand=True)

        # Frame per il testo informativo con sfondo bianco e bordo
        info_frame = tk.Frame(body_frame, bg="#FFFFFF", relief="solid", borderwidth=2, padx=15, pady=15)
        info_frame.pack(fill="both", expand=True, pady=(0, 20))

        info_label = tk.Label(info_frame,
                             text="Il Livello di Rischio Effettivo calcolato supera la soglia di 3.5,\nconfigurano un caso di RISCHIO RAFFORZATO.\n\n" +
                                  "Ai sensi del D.Lgs. 231/2007 (Artt. 17, 23, 28) e delle Linee Guida VEDA\n" +
                                  "Edizione 02-2020, il Professionista è OBBLIGATO ad applicare\n" +
                                  "l'Adeguata Verifica Rafforzata, che include tra gli altri obblighi:\n\n" +
                                  "• Acquisizione e analisi documentale della PROVENIENZA DEI FONDI\n" +
                                  "• Acquisizione e analisi documentale della DESTINAZIONE DEI FONDI\n" +
                                  "• Verifica rafforzata dell'identità del cliente e del titolare effettivo\n" +
                                  "• Acquisizione di informazioni sullo scopo e sulla natura del rapporto\n" +
                                  "• Monitoraggio continuo del rapporto professionale\n" +
                                  "• Conservazione della documentazione per almeno 10 anni\n\n" +
                                  "L'omissione di tali obblighi può comportare sanzioni amministrative\n" +
                                  "pecuniarie (da €2.500 a €50.000 per le violazioni più gravi)\n" +
                                  "e responsabilità disciplinare.",
                             font=("Helvetica", 10), bg="#FFFFFF", fg="#000000",
                             justify="left", anchor="nw")
        info_label.pack(fill="both", expand=True)

        # Checkbox obbligatoria
        checkbox_var = tk.BooleanVar(value=False)
        checkbox_frame = tk.Frame(body_frame, bg="#FFF3CD")
        checkbox_frame.pack(fill="x", pady=(0, 20))

        checkbox = tk.Checkbutton(checkbox_frame,
                                 text="Dichiaro di aver preso visione degli obblighi di Adeguata Verifica Rafforzata e mi impegno ad adempierli",
                                 variable=checkbox_var, bg="#FFF3CD", font=("Helvetica", 10, "bold"),
                                 fg="#8B0000", activebackground="#FFF3CD", wraplength=700, justify="left")
        checkbox.pack(anchor="w")

        # Risultato (None = annulla, Dict = accettato)
        risultato = {"value": None}

        def conferma():
            if not checkbox_var.get():
                messagebox.showwarning("Attenzione", "Devi spuntare la casella di accettazione per procedere.", parent=dialog)
                return

            from datetime import datetime
            timestamp = datetime.now().strftime("%d/%m/%Y alle ore %H:%M:%S")
            risultato["value"] = {
                "accepted": True,
                "timestamp": timestamp,
                "somma_ponderata": somma_ponderata,
                "livello": livello
            }
            dialog.destroy()

        def annulla():
            risposta = messagebox.askyesno("Conferma Annullamento",
                                          "Sei sicuro di voler annullare?\nLa valutazione NON verrà salvata.",
                                          parent=dialog)
            if risposta:
                risultato["value"] = None
                dialog.destroy()

        # Bottoni
        button_frame = tk.Frame(body_frame, bg="#FFF3CD")
        button_frame.pack(fill="x")

        btn_conferma = tk.Button(button_frame, text="✓ Confermo e Procedo", command=conferma,
                                bg="#28A745", fg="white", font=("Helvetica", 12, "bold"),
                                padx=20, pady=10)
        btn_conferma.pack(side="left", expand=True, padx=(0, 10))

        btn_annulla = tk.Button(button_frame, text="✗ Annulla Valutazione", command=annulla,
                               bg="#DC3545", fg="white", font=("Helvetica", 12, "bold"),
                               padx=20, pady=10)
        btn_annulla.pack(side="left", expand=True, padx=(10, 0))

        # Impedisce chiusura con X senza conferma
        def on_closing():
            annulla()

        dialog.protocol("WM_DELETE_WINDOW", on_closing)

        # Aspetta che il dialog venga chiuso
        dialog.wait_window()

        return risultato["value"]

    def esporta_word(self) -> None:
        if self.dati_export is None:
            messagebox.showerror("Errore", "Prima esegui la valutazione del rischio!")
            return
        try:
            sections_a = Config.get_sections_A_manual()
            sections_b = Config.get_sections_B_manual()
            WordExporter.esporta(self.dati_export, sections_a, sections_b)
        except Exception as e:
            messagebox.showerror("Errore", f"Errore durante l'export Word:\n{str(e)}")

def main():
    """Entry point dell'applicazione"""
    Config.inizializza_luoghi()
    root = tk.Tk()
    app = AMLRiskApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()